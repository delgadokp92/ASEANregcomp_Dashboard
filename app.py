import base64
import io
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from urllib.parse import urlparse

import pandas as pd
import plotly.express as px
import streamlit as st
import streamlit.components.v1 as components

# =========================
# App config
# =========================
st.set_page_config(
    page_title="ARIS",
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items={
        "About": "ARIS — ASEAN Regulatory Information System: interactive view of regional regulations.",
    },
)

st.markdown(
    """
    <style>
    /* ── Base ── */
    .stApp { background-color: #071014; color: #f8fafc; }

    /* ── Collapse the Streamlit header to zero height ── */
    header[data-testid="stHeader"] {
        height: 0 !important; min-height: 0 !important;
        background: transparent !important; border-bottom: none !important;
        pointer-events: none !important; overflow: hidden !important;
    }

    /* ── Hide sidebar and its toggle permanently ── */
    [data-testid="stSidebar"],
    [data-testid="stSidebarCollapseButton"],
    [data-testid="collapsedControl"] {
        display: none !important;
    }

    /* ── Full-width main content, flush to top ── */
    section[data-testid="stMain"] { margin-left: 0 !important; width: 100% !important; }
    .block-container { padding-top: 0.5rem !important; }

    /* ── Top nav bar ── */
    .nav-bar {
        display: flex; align-items: center; gap: 6px;
        background: #0d1526; border-bottom: 1px solid #1e293b;
        padding: 6px 12px; margin: -0.5rem -1rem 0.75rem -1rem;
        position: sticky; top: 0; z-index: 100;
    }
    .nav-title {
        font-size: 0.95rem; font-weight: 700; color: #e2e8f0;
        white-space: nowrap; margin-right: 8px; letter-spacing: 0.01em;
    }
    .nav-version {
        font-size: 0.62rem; color: #475569; white-space: nowrap; margin-right: 12px;
    }

    /* ── Native dataframe — match custom HTML table palette ── */
    /* Outer border and radius match the iframe tables */
    [data-testid="stDataFrame"] {
        border: 1px solid #334155 !important;
        border-radius: 6px !important;
        overflow: hidden !important;
    }
    [data-testid="stDataFrame"] > div { border-radius: 6px !important; }
    /* Scrollbar track — match dark bg */
    [data-testid="stDataFrame"] ::-webkit-scrollbar { width: 6px; height: 6px; }
    [data-testid="stDataFrame"] ::-webkit-scrollbar-track { background: #0d1526; }
    [data-testid="stDataFrame"] ::-webkit-scrollbar-thumb {
        background: #334155; border-radius: 3px;
    }
    [data-testid="stDataFrame"] ::-webkit-scrollbar-thumb:hover { background: #475569; }

    /* ── Force dark theme on all Streamlit controls (overrides OS light-mode default) ── */
    /* Secondary buttons */
    [data-testid="stBaseButton-secondary"] {
        background-color: #1e293b !important;
        color: #e2e8f0 !important;
        border: 1px solid #334155 !important;
    }
    [data-testid="stBaseButton-secondary"]:hover {
        background-color: #273548 !important;
        border-color: #475569 !important;
        color: #f1f5f9 !important;
    }
    /* Primary buttons */
    [data-testid="stBaseButton-primary"] {
        background-color: #3b82f6 !important;
        color: #ffffff !important;
        border: none !important;
    }
    [data-testid="stBaseButton-primary"]:hover {
        background-color: #2563eb !important;
    }
    /* Download buttons */
    [data-testid="stDownloadButton"] > button {
        background-color: #1e293b !important;
        color: #e2e8f0 !important;
        border: 1px solid #334155 !important;
    }
    [data-testid="stDownloadButton"] > button:hover {
        background-color: #273548 !important;
        border-color: #475569 !important;
    }
    /* Selectbox, text input, text area */
    [data-testid="stSelectbox"] > div > div,
    [data-baseweb="select"] > div,
    [data-testid="stTextInput"] > div > div > input,
    [data-testid="stTextArea"] > div > div > textarea {
        background-color: #1e293b !important;
        color: #e2e8f0 !important;
        border-color: #334155 !important;
    }
    /* Dropdown menu list */
    [data-baseweb="popover"] ul,
    [data-baseweb="menu"] {
        background-color: #1e293b !important;
        color: #e2e8f0 !important;
    }
    [data-baseweb="menu"] li:hover {
        background-color: #273548 !important;
    }
    /* Expander headers */
    [data-testid="stExpander"] summary {
        background-color: #0d1526 !important;
        color: #e2e8f0 !important;
    }
    /* Tabs (st.tabs) */
    [data-testid="stTabs"] [data-baseweb="tab-list"] {
        background-color: #0f172a !important;
    }
    [data-testid="stTabs"] [data-baseweb="tab"] {
        color: #94a3b8 !important;
        background-color: transparent !important;
    }
    [data-testid="stTabs"] [aria-selected="true"] {
        color: #e2e8f0 !important;
        border-bottom-color: #3b82f6 !important;
    }
    /* Popover panel */
    [data-testid="stPopover"] > div {
        background-color: #1e293b !important;
        border: 1px solid #334155 !important;
        color: #e2e8f0 !important;
    }
    /* Dialog / modal */
    [data-testid="stModal"] > div > div {
        background-color: #0f172a !important;
        color: #e2e8f0 !important;
    }

    /* ── Typography ── */
    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4 {
        color: #e2e8f0; letter-spacing: 0.01em;
    }

    /* ── Regulation table links ── */
    .regulations-table a {
        color: #7dd3fc; text-decoration: none; transition: color 0.15s ease;
    }
    .regulations-table a:hover { color: #bae6fd; text-decoration: underline; }

    /* ── Mobile ── */
    @media (max-width: 768px) {
        .stApp { font-size: 14px; }
        h1 { font-size: 1.4rem !important; }
        h2 { font-size: 1.2rem !important; }
        h3 { font-size: 1.05rem !important; }
        .regulations-table td {
            display: block; width: 100% !important;
            text-align: left !important; padding: 4px 0 !important;
        }
    }

    /* ── Print / Save as PDF — landscape, white bg, black fonts, map preserved ── */
    @page { size: A4 landscape; margin: 1.2cm; }

    @media print {
        /* White backgrounds on Streamlit containers — not on SVG elements */
        html, body, .stApp,
        section[data-testid="stMain"],
        div[data-testid="stMainBlockContainer"],
        .block-container,
        div[data-testid="stVerticalBlock"],
        div[data-testid="stHorizontalBlock"],
        div[data-testid="stColumn"],
        div[data-testid="element-container"],
        div[data-baseweb="tab-panel"],
        div[data-testid="stExpander"],
        div[data-testid="stPlotlyChart"],
        div[data-testid="stPlotlyChart"] > div,
        div[class*="st-emotion-cache"] {
            background: #ffffff !important;
            background-color: #ffffff !important;
            box-shadow: none !important;
            border-color: #cccccc !important;
        }

        /* Black text universally — CSS `color` does NOT affect SVG fill so map stays.
           opacity:1 removes Streamlit's muted/secondary element fading. */
        * { color: #000000 !important; opacity: 1 !important; }
        a, a:visited { color: #0055aa !important; }

        /* Plotly SVG annotation/axis text → black */
        [data-testid="stPlotlyChart"] svg text,
        [data-testid="stPlotlyChart"] svg tspan { fill: #000000 !important; }

        /* Map: stretch to full page width, left-aligned, preserve flex layout */
        [data-testid="stPlotlyChart"] {
            width: 100% !important;
            max-width: 100% !important;
            margin-left: 0 !important;
        }
        [data-testid="stPlotlyChart"] .js-plotly-plot,
        [data-testid="stPlotlyChart"] .plot-container { width: 100% !important; }

        /* Hide interactive chrome */
        [data-testid="stSidebar"],
        header[data-testid="stHeader"],
        [data-testid="stToolbar"],
        button,
        [data-testid="stBaseButton-secondary"],
        [data-testid="stBaseButton-primary"],
        [data-testid="stDownloadButton"],
        [data-testid="stExpanderToggleIcon"],
        [data-testid="stDecoration"],
        .modebar, .modebar-container { display: none !important; }

        /* Page breaks */
        h2, h3 { page-break-after: avoid; }
        table, [data-testid="stPlotlyChart"] { page-break-inside: avoid; }
    }
    </style>
    """,
    unsafe_allow_html=True,
)
DATA_DIR = Path("src") / "categories"
ARCHIVE_FILE = Path("src") / "CBregs_audit_log.csv"


def resolve_data_dir(path: Path) -> Path:
    candidates = [
        path,
        Path(__file__).resolve().parent / path,
    ]
    for candidate in candidates:
        if candidate.is_dir():
            return candidate
    raise FileNotFoundError(
        f"Data directory not found. Tried: {', '.join(str(p) for p in candidates)}"
    )


def get_csv_cache_key(src_dir: Path) -> tuple:
    csv_files = sorted(src_dir.glob("*.csv"))
    return tuple((f.name, f.stat().st_mtime) for f in csv_files)


# =========================
# Helpers
# =========================
ASEAN_FLAG = {
    "Brunei Darussalam": "🇧🇳",
    "Cambodia": "🇰🇭",
    "Indonesia": "🇮🇩",
    "Lao PDR": "🇱🇦",
    "Laos": "🇱🇦",
    "Malaysia": "🇲🇾",
    "Myanmar": "🇲🇲",
    "Philippines": "🇵🇭",
    "Singapore": "🇸🇬",
    "Thailand": "🇹🇭",
    "Viet Nam": "🇻🇳",
    "Vietnam": "🇻🇳",
    "Timor-Leste": "🇹🇱",
}

ASEAN_ISO3 = {
    "Brunei Darussalam": "BRN",
    "Cambodia":          "KHM",
    "Indonesia":         "IDN",
    "Lao PDR":           "LAO",
    "Laos":              "LAO",
    "Malaysia":          "MYS",
    "Myanmar":           "MMR",
    "Philippines":       "PHL",
    "Singapore":         "SGP",
    "Thailand":          "THA",
    "Viet Nam":          "VNM",
    "Vietnam":           "VNM",
    "Timor-Leste":       "TLS",
}


def _html_src(html: str) -> str:
    """Encode HTML as a base64 data URL for st.iframe."""
    b64 = base64.b64encode(html.encode("utf-8")).decode("ascii")
    return f"data:text/html;base64,{b64}"

COUNTRY_ISO_CODES = {
    "Brunei Darussalam": "BN",
    "Cambodia": "KH",
    "Indonesia": "ID",
    "Lao PDR": "LA",
    "Laos": "LA",
    "Malaysia": "MY",
    "Myanmar": "MM",
    "Philippines": "PH",
    "Singapore": "SG",
    "Thailand": "TH",
    "Viet Nam": "VN",
    "Vietnam": "VN",
    "Timor-Leste": "TL",
}


_APP_VERSION = "3.9"
_APP_DATE    = "2026-06-30"

# Design tokens (use in inline styles for consistency)
_C_BG_DEEP  = "#071014"; _C_BG_BASE  = "#0f172a"; _C_BG_CARD  = "#0d1526"
_C_BG_ROW1  = "#111e33"; _C_BG_PANEL = "#1e293b"; _C_BORDER   = "#334155"
_C_TEXT     = "#e2e8f0"; _C_MUTED    = "#94a3b8"; _C_FAINT    = "#475569"
_C_LINK     = "#93c5fd"; _C_GREEN    = "#22c55e"; _C_AMBER    = "#f59e0b"
_C_RED      = "#ef4444"; _C_GRAY     = "#64748b"


def country_code_to_emoji(code: str) -> str:
    code = str(code).upper().strip()
    if len(code) != 2 or not code.isalpha():
        return ""
    return "".join(chr(ord(ch) + 0x1F1E6 - ord("A")) for ch in code)


def country_flag(country: str) -> str:
    country = str(country).strip()
    if not country:
        return "🏳️"
    if country in ASEAN_FLAG:
        return ASEAN_FLAG[country]
    iso = COUNTRY_ISO_CODES.get(country)
    if iso:
        return country_code_to_emoji(iso)
    return "🏳️"


META_COL_CANDIDATES = {
    "country": ["Country_std", "country", "Country"],
    "regulator": ["Regulator_std", "regulator", "Regulator"],
    "year": [
        "Year_raw",
        "year",
        "Year",
        "Year approved/implemented",
        "Year Approved/Implemented",
        "Year approved / implemented",
    ],
    "source": ["Source_Original", "Source_URL", "source", "Official Source", "Official source", "Official Source links", "Official source links", "Source", "URL", "Link"],
    "title": [
        "Issuance",
        "Regulation_Title",
        "title",
        "Regulation / Legal Instrument",
        "Regulation / Legal instrument",
        "Primary Legal / Regulatory Framework",
        "Primary Legal/Regulatory Framework",
        "Regulations on fraud risk management",
        "Regulations on consumer protection (payments)",
        "Regulation",
        "Legal Instrument",
    ],
}

_KNOWN_META: set[str] = set(
    META_COL_CANDIDATES["country"]
    + META_COL_CANDIDATES["regulator"]
    + META_COL_CANDIDATES["year"]
    + META_COL_CANDIDATES["source"]
    + META_COL_CANDIDATES["title"]
    + [
        "Category", "Country_std", "Regulator_std", "Year_raw", "Year",
        "Year_sort", "Regulation_Title", "Source_URL", "Entry_ID", "HasData", "ID",
        "Source_Original", "Source_EN", "Amendment_Of",
    ]
)


def get_provision_cols(df: pd.DataFrame, category: str) -> List[str]:
    cat_df = df[df["Category"] == category]
    return [
        c for c in cat_df.columns
        if c not in _KNOWN_META and not cat_df[c].isna().all()
    ]


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [re.sub(r"\s+", " ", str(c)).strip() for c in df.columns]
    return df


def extract_year(value) -> Optional[int]:
    if pd.isna(value):
        return None
    s = str(value).strip()
    # Strip effectivity date qualifiers so "2019 (effective 2022)" → 2019
    s = re.sub(r'\(effectiv\w*[^)]*\)', '', s, flags=re.IGNORECASE)
    s = re.sub(r'effectiv\w*\s+\d[^,;|]*', '', s, flags=re.IGNORECASE)
    m = re.search(r"(19\d{2}|20\d{2})", s)
    return int(m.group(1)) if m else None


def format_year(yr) -> str:
    if yr is None or (isinstance(yr, float) and pd.isna(yr)):
        return "—"
    s = str(yr).strip()
    if s.lower() in ("nan", "none", ""):
        return "—"
    try:
        return str(int(float(s)))
    except (ValueError, OverflowError):
        return "—"


def pick_first_existing_col(
    df: pd.DataFrame, candidates: List[str]
) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    return None


def infer_title_col(df: pd.DataFrame) -> Optional[str]:
    # 1) Try known title candidates.
    c = pick_first_existing_col(df, META_COL_CANDIDATES["title"])
    if c:
        return c

    # 2) Fallback: pick the first non-meta column whose values look like text.
    known_meta = set(
        META_COL_CANDIDATES["country"]
        + META_COL_CANDIDATES["regulator"]
        + META_COL_CANDIDATES["year"]
        + META_COL_CANDIDATES["source"]
        + ["Regulation ID"]
    )
    for col in df.columns:
        if col in known_meta:
            continue
        if df[col].astype(str).str.len().mean() > 5:
            return col
    return None


def safe_linkify(url) -> str:
    if url is None:
        return ""
    url = str(url).strip()
    if not url or url.lower() in {"nan", "none"}:
        return ""
    parsed = urlparse(url)
    if not parsed.scheme:
        url = f"https://{url}"
    return url


def html_escape(text: str) -> str:
    text = str(text)
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
            .replace("\n", " ")
            .replace("\r", " ")
    )


def _render_provision_table(rows: list[dict], columns: list[str]) -> None:
    """Render a provision table as a fixed-height iframe (scrollable if content overflows)."""
    import html as _hl

    # Build <th> cells — each has a drag handle div
    th_cells = "".join(
        f'<th><span class="col-label">{_hl.escape(str(c))}</span>'
        f'<div class="resizer" title="Drag to resize"></div></th>'
        for c in columns
    )

    def _cell_html(raw: str) -> str:
        """Escape HTML but convert [text](url) markdown links to <a> tags first."""
        parts: list[str] = []
        last = 0
        for m in re.finditer(r'\[([^\]]+)\]\(([^)]+)\)', raw):
            parts.append(_hl.escape(raw[last:m.start()]))
            txt = _hl.escape(m.group(1))
            url = _hl.escape(m.group(2).strip())
            parts.append(
                f'<a href="{url}" target="_blank" rel="noopener"'
                f' style="color:#7dd3fc;text-decoration:underline">{txt}</a>'
            )
            last = m.end()
        parts.append(_hl.escape(raw[last:]))
        return "".join(parts).replace("\r\n", "<br>").replace("\n", "<br>")

    # Build <tr> cells
    body_rows = []
    for i, row in enumerate(rows):
        cells = []
        for j, col in enumerate(columns):
            raw = str(row.get(col, ""))
            val = _cell_html(raw)
            cls = "fc" if j == 0 else "vc"
            cells.append(f'<td class="{cls}">{val}</td>')
        cls = "r0" if i % 2 == 0 else "r1"
        body_rows.append(f'<tr class="{cls}">{"".join(cells)}</tr>')

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:transparent;font-family:"Source Sans Pro",sans-serif;font-size:14px;color:#e2e8f0}}
table{{width:100%;border-collapse:collapse;table-layout:fixed}}
thead tr{{background:#1e293b}}
th{{position:relative;padding:7px 10px 7px 8px;text-align:left;color:#94a3b8;
    font-size:13px;font-weight:600;border-bottom:2px solid #334155;
    white-space:normal;overflow:visible;user-select:none;vertical-align:bottom}}
.col-label{{display:block;padding-right:6px;line-height:1.35;word-break:break-word}}
.resizer{{position:absolute;right:0;top:0;bottom:0;width:5px;cursor:col-resize;
          background:transparent;z-index:1}}
.resizer:hover,.resizer.active{{background:#3b82f6;opacity:.7}}
tr.r0{{background:#0d1526}}
tr.r1{{background:#111e33}}
tr:hover{{background:#1a2d4a}}
td{{padding:6px 10px;vertical-align:top;border-bottom:1px solid #1e293b;
    line-height:1.6;word-break:break-word;overflow-wrap:anywhere;font-size:13.5px}}
td.fc{{color:#7dd3fc;font-size:13px;font-weight:600;word-break:break-word;overflow-wrap:anywhere}}
td.vc{{font-size:13.5px}}
@media (max-width:600px){{td,th{{font-size:11px!important;padding:4px 6px!important}}}}
@media print{{body,table,thead tr,tr.r0,tr.r1,td,th{{background:#fff!important;color:#111!important;border-color:#ccc!important}}a{{color:#0055aa!important}}}}
</style>
</head>
<body>
<table id="t">
  <thead><tr>{th_cells}</tr></thead>
  <tbody>{"".join(body_rows)}</tbody>
</table>
<script>
(function(){{
  const tbl = document.getElementById('t');
  const ths = tbl.querySelectorAll('th');

  // Set initial column widths
  const total = tbl.parentElement.offsetWidth || 800;
  const fieldW = Math.min(170, total * 0.18);
  const otherW = (total - fieldW) / Math.max(1, ths.length - 1);
  ths[0].style.width = fieldW + 'px';
  for (let i = 1; i < ths.length; i++) ths[i].style.width = otherW + 'px';

  // Resize logic
  let active = null, startX = 0, startW = 0;
  tbl.querySelectorAll('.resizer').forEach(r => {{
    r.addEventListener('mousedown', e => {{
      active = r;
      startX = e.pageX;
      startW = r.parentElement.offsetWidth;
      r.classList.add('active');
      document.body.style.cursor = 'col-resize';
      e.preventDefault();
    }});
  }});
  document.addEventListener('mousemove', e => {{
    if (!active) return;
    const w = Math.max(60, startW + (e.pageX - startX));
    active.parentElement.style.width = w + 'px';
  }});
  document.addEventListener('mouseup', () => {{
    if (active) {{
      active.classList.remove('active');
      active = null;
      document.body.style.cursor = '';
    }}
  }});

  // Report table height to Streamlit so iframe fits exactly
  let _lastH = 0;
  function sendHeight() {{
    const h = Math.ceil(tbl.getBoundingClientRect().height) + 8;
    if (h > 0 && h !== _lastH) {{ _lastH = h;
      window.parent.postMessage({{isStreamlitMessage:true,type:'streamlit:setFrameHeight',height:h}}, '*');
    }}
  }}
  sendHeight();
  window.addEventListener('load', sendHeight);
  setTimeout(sendHeight, 100);
  setTimeout(sendHeight, 500);
  new ResizeObserver(sendHeight).observe(tbl);
}})();
</script>
</body></html>"""

    total_chars = sum(len(str(row.get(col, ""))) for row in rows for col in columns[1:])
    est = max(300, 60 + len(rows) * 60 + (total_chars // 60) * 22)
    st.iframe(src=_html_src(html), height=est)


def _render_data_table(df: pd.DataFrame, highlighted: set | None = None, row_height: int = 72) -> None:
    """Render a DataFrame as a resizable HTML table; newlines in cells become <br>.
    First column (flag) is centred; second column (country) is bold.
    Rows whose Country value is in `highlighted` are tinted."""
    import html as _hl

    CHECK = "✓"
    # Detect which columns (beyond Flag/Country) contain only checkmarks or blanks
    check_cols: set[str] = set()
    for col in list(df.columns)[2:]:
        vals = df[col].fillna("").astype(str).str.strip()
        if vals.isin({CHECK, ""}).all():
            check_cols.add(col)

    def _col_min_width(name: str, is_check: bool) -> str:
        # Base min-width on the longest word so no word is ever mid-truncated.
        # ~0.62rem per character is a rough but reliable estimate at 13px font.
        longest = max((len(w) for w in str(name).split()), default=4)
        rems = longest * 0.62
        floor = 4.0 if is_check else 8.0
        return f"{max(floor, rems):.1f}rem"

    def _th(c: str) -> str:
        label = _hl.escape(str(c))
        if str(c) == "Flag":
            return (
                '<th style="width:2rem;min-width:2rem;padding:7px 4px 7px 8px;text-align:center;">'
                '<span class="col-label"></span>'
                '<div class="resizer" title="Drag to resize"></div></th>'
            )
        if str(c) in check_cols:
            mw = _col_min_width(c, is_check=True)
            return (
                f'<th style="min-width:{mw};text-align:center;">'
                f'<span class="col-label" style="text-align:center;">{label}</span>'
                f'<div class="resizer" title="Drag to resize"></div></th>'
            )
        mw = _col_min_width(c, is_check=False)
        return (
            f'<th style="min-width:{mw};">'
            f'<span class="col-label">{label}</span>'
            f'<div class="resizer" title="Drag to resize"></div></th>'
        )

    th_cells = "".join(_th(c) for c in df.columns)

    body_rows = []
    for i, (_, row) in enumerate(df.iterrows()):
        country_val = str(row.get("Country", "")) if "Country" in df.columns else ""
        is_sel = bool(highlighted and country_val in highlighted)
        cells = []
        for j, col in enumerate(df.columns):
            raw = "" if pd.isna(row[col]) else str(row[col])
            val = _hl.escape(raw).replace("\r\n", "<br>").replace("\n", "<br>")
            if j == 0:
                cells.append(f'<td style="text-align:center;padding:6px 4px 6px 8px;vertical-align:middle;white-space:nowrap;width:2rem;">{val}</td>')
            elif j == 1:
                cells.append(f'<td style="font-weight:600;padding:6px 10px;vertical-align:top;white-space:nowrap;">{val}</td>')
            elif col in check_cols:
                cells.append(f'<td class="vc chk">{val}</td>')
            else:
                cells.append(f'<td class="vc">{val}</td>')
        cls = "sel" if is_sel else ("r0" if i % 2 == 0 else "r1")
        body_rows.append(f'<tr class="{cls}">{"".join(cells)}</tr>')

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:transparent;font-family:"Source Sans Pro",sans-serif;font-size:14px;color:#e2e8f0}}
table{{border-collapse:collapse;table-layout:auto;width:100%}}
thead tr{{background:#1e293b}}
th{{position:relative;padding:7px 10px 7px 8px;text-align:left;color:#94a3b8;
    font-size:13px;font-weight:600;border-bottom:2px solid #334155;
    white-space:normal;overflow:visible;user-select:none;vertical-align:bottom}}
.col-label{{display:block;padding-right:6px;line-height:1.35;word-break:break-word;hyphens:auto}}
.resizer{{position:absolute;right:0;top:0;bottom:0;width:5px;cursor:col-resize;
          background:transparent;z-index:1}}
.resizer:hover,.resizer.active{{background:#3b82f6;opacity:.7}}
tr.r0{{background:#0d1526}}
tr.r1{{background:#111e33}}
tr.sel{{background:#1a3a5c}}
tr:hover{{background:#1a2d4a}}
tr.sel:hover{{background:#20446b}}
td{{padding:6px 10px;vertical-align:top;border-bottom:1px solid #1e293b;
    line-height:1.6;word-break:break-word;overflow-wrap:anywhere;font-size:13.5px}}
.vc{{font-size:13.5px}}
.chk{{text-align:center;color:#4ade80;font-size:15px;vertical-align:middle;}}
@media (max-width:600px){{td,th{{font-size:11px!important;padding:4px 6px!important}}}}
@media print{{body,table,thead tr,tr.r0,tr.r1,td,th{{background:#fff!important;color:#111!important;border-color:#ccc!important}}a{{color:#0055aa!important}}}}
</style>
</head>
<body>
<table id="t">
  <thead><tr>{th_cells}</tr></thead>
  <tbody>{"".join(body_rows)}</tbody>
</table>
<script>
(function(){{
  const tbl = document.getElementById('t');

  // Phase 1: let auto-layout determine natural widths from content/min-width
  tbl.style.tableLayout = 'auto';

  function initResize() {{
    // Phase 2: snapshot actual widths, switch to fixed so drag-resize works
    const ths = Array.from(tbl.querySelectorAll('th'));
    ths.forEach(th => {{ th.style.width = th.offsetWidth + 'px'; }});
    tbl.style.tableLayout = 'fixed';
    tbl.style.width = tbl.offsetWidth + 'px';

    // Wire up drag-to-resize on each column
    let active = null, startX = 0, startW = 0;
    ths.forEach(th => {{
      const r = th.querySelector('.resizer');
      if (!r) return;
      r.addEventListener('mousedown', e => {{
        active = th; startX = e.pageX; startW = th.offsetWidth;
        r.classList.add('active');
        document.body.style.cursor = 'col-resize';
        e.preventDefault();
      }});
    }});
    document.addEventListener('mousemove', e => {{
      if (!active) return;
      active.style.width = Math.max(40, startW + (e.pageX - startX)) + 'px';
    }});
    document.addEventListener('mouseup', () => {{
      if (active) {{
        active.querySelector('.resizer')?.classList.remove('active');
        active = null;
        document.body.style.cursor = '';
        sendH();
      }}
    }});
  }}

  let _lastH = 0;
  function sendH() {{
    const h = Math.ceil(tbl.getBoundingClientRect().height) + 8;
    if (h > 0 && h !== _lastH) {{
      _lastH = h;
      window.parent.postMessage({{isStreamlitMessage:true,type:'streamlit:setFrameHeight',height:h}},'*');
    }}
  }}

  // Small delay lets the browser complete auto-layout before we snapshot
  setTimeout(() => {{ initResize(); sendH(); }}, 80);
  window.addEventListener('load', sendH);
  new ResizeObserver(sendH).observe(tbl);
}})();
</script>
</body></html>"""

    est = min(900, max(150, 44 + len(df) * row_height))
    st.iframe(src=_html_src(html), height=est)


def build_regulations_html_table(df: pd.DataFrame) -> str:
    rows = [
        '<table class="regulations-table" style="width:100%; border-collapse: collapse; margin-bottom:1rem; table-layout: fixed;">',
        "<tbody>",
    ]
    for row in df.itertuples(index=False):
        title = getattr(row, "Regulation_Title", None)
        title = html_escape(title) if pd.notna(title) else "—"
        year = getattr(row, "Year", None)
        year_text = format_year(year)
        source_url = safe_linkify(getattr(row, "Source_URL", None))
        _source_en_raw = getattr(row, "Source_EN", None)
        source_en_url = safe_linkify(_source_en_raw) if pd.notna(_source_en_raw) else ""
        if source_url and source_en_url and source_en_url != source_url:
            source_html = (
                f'<a href="{html_escape(source_url)}" target="_blank" rel="noreferrer">Source</a>'
                f' · <a href="{html_escape(source_en_url)}" target="_blank" rel="noreferrer">EN</a>'
            )
        elif source_url:
            source_html = (
                f'<a href="{html_escape(source_url)}" target="_blank" rel="noreferrer">Source</a>'
            )
        else:
            source_html = "—"
        rows.append(
            "<tr>"
            f"<td style=\"padding:8px 12px 8px 0; vertical-align:top; white-space:normal; word-break:break-word;\">{title}</td>"
            f"<td style=\"padding:8px 12px; vertical-align:top; white-space:nowrap; text-align:right;\">{year_text}</td>"
            f"<td style=\"padding:8px 0 8px 12px; vertical-align:top; white-space:nowrap; text-align:center;\">{source_html}</td>"
            "</tr>"
        )
    rows.extend(["</tbody>", "</table>"])
    return "\n".join(rows)


def get_selected_row_indices(event) -> List[int]:
    try:
        return list(event.selection.rows) if event and event.selection else []
    except Exception:
        return []


@st.cache_data
def load_cbregs(src_dir: str, cache_key: tuple) -> pd.DataFrame:
    p = Path(src_dir)
    csv_files = sorted(p.glob("*.csv"))
    frames = []

    for csv_file in csv_files:
        sheet = csv_file.stem
        df = pd.read_csv(csv_file, dtype=object)
        df = normalize_columns(df).dropna(how="all").dropna(axis=1, how="all")

        country_col = pick_first_existing_col(df, META_COL_CANDIDATES["country"])
        regulator_col = pick_first_existing_col(df, META_COL_CANDIDATES["regulator"])
        year_col = pick_first_existing_col(df, META_COL_CANDIDATES["year"])
        source_col = pick_first_existing_col(df, META_COL_CANDIDATES["source"])
        title_col = infer_title_col(df)

        # Build standardized view while retaining originals for the country modal
        out = df.copy()
        out["Category"] = sheet

        out["Country_std"] = out[country_col] if country_col else pd.NA
        out["Regulator_std"] = out[regulator_col] if regulator_col else pd.NA
        out["Year_raw"] = out[year_col] if year_col else pd.NA
        out["Year"] = out["Year_raw"].apply(extract_year).astype("Int64")

        if title_col:
            # Preserve any pre-existing Regulation_Title (e.g. rows added via Editor)
            # before overwriting with the sheet's native title column.
            _pre_rt = (
                out["Regulation_Title"].copy()
                if "Regulation_Title" in out.columns and title_col != "Regulation_Title"
                else None
            )
            out["Regulation_Title"] = out[title_col].astype(str)
            if _pre_rt is not None:
                _blank = out["Regulation_Title"].str.strip().isin(["", "nan", "None", "<NA>"])
                out.loc[_blank, "Regulation_Title"] = _pre_rt[_blank].astype(str)
        else:
            out["Regulation_Title"] = pd.NA

        if source_col:
            out["Source_URL"] = out[source_col].astype(str)
        else:
            out["Source_URL"] = pd.NA

        # Source_EN — English translation URL (same as Source_URL if not available)
        if "Source_EN" in df.columns:
            out["Source_EN"] = out["Source_EN"].astype(str)
        else:
            out["Source_EN"] = pd.NA

        # Amendment_Of — reference to original regulation for amendments
        if "Amendment_Of" in df.columns:
            out["Amendment_Of"] = out["Amendment_Of"].astype(str)
        else:
            out["Amendment_Of"] = pd.NA

        frames.append(out)

    all_df = pd.concat(frames, ignore_index=True)

    # Clean
    all_df["Country_std"] = all_df["Country_std"].astype(str).str.strip()
    all_df["Regulator_std"] = all_df["Regulator_std"].astype(str).str.strip()
    all_df["Regulation_Title"] = all_df["Regulation_Title"].astype(str).str.strip()

    # Treat "nan" strings produced by astype(str)
    for c in ["Country_std", "Regulator_std", "Regulation_Title", "Source_URL", "Source_EN", "Amendment_Of"]:
        if c in all_df.columns:
            all_df.loc[all_df[c].str.lower().isin(["nan", "none", ""]), c] = pd.NA

    all_df = ensure_entry_ids(all_df)
    return all_df


def latest_regs_by_country(df: pd.DataFrame, country: str, n: int = 10) -> pd.DataFrame:
    d = df[df["Country_std"] == country].copy()
    d = d.dropna(subset=["Regulation_Title"])
    # Year might be NA; put those last
    d["Year_sort"] = d["Year"].fillna(-1).astype(int)
    d = d.sort_values(["Year_sort", "Regulation_Title"], ascending=[False, True])
    return d.head(n)[["Year", "Regulation_Title", "Category", "Regulator_std", "Source_URL"]]


def build_hover_list(df_country_latest: pd.DataFrame) -> str:
    if df_country_latest.empty:
        return "No regulations found."
    lines = []
    for _, r in df_country_latest.iterrows():
        y = r["Year"]
        y_txt = str(int(y)) if pd.notna(y) else "—"
        title = str(r["Regulation_Title"])
        if len(title) > 60:
            title = title[:57] + "…"
        lines.append(f"{y_txt} — {title}")
    # Plotly hover supports <br>
    return "<br>".join(lines)


def ensure_entry_ids(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    if "Entry_ID" not in df.columns:
        df["Entry_ID"] = pd.NA

    df["Entry_ID"] = df["Entry_ID"].astype(str).replace(["nan", "None", "none"], pd.NA)
    missing = df["Entry_ID"].isna() | (df["Entry_ID"] == "")
    if missing.any():
        df.loc[missing, "Entry_ID"] = [str(uuid.uuid4()) for _ in range(missing.sum())]

    return df


def serialize_value_for_archive(value):
    if pd.isna(value):
        return None
    if isinstance(value, (dict, list)):
        return value
    return str(value)


def serialize_record_for_archive(record: pd.Series) -> dict:
    return {
        k: serialize_value_for_archive(v)
        for k, v in record.items()
        if k is not None
    }


_AUDIT_KEY_FIELDS = [
    "Regulation_Title", "Year", "Category", "Regulator_std",
    "Amendment_Of", "Source_URL", "Source_EN",
]

def _audit_changes(action: str, old_str: str, new_str: str) -> str:
    """Return a human-readable summary of what changed in an audit row."""
    try:
        old = json.loads(old_str) if old_str else {}
        new = json.loads(new_str) if new_str else {}
    except Exception:
        return ""

    def _trunc(v: str, n: int = 40) -> str:
        v = str(v or "").strip()
        return (v[:n] + "…") if len(v) > n else v

    if action == "add":
        title = _trunc(new.get("Regulation_Title", ""), 50)
        cat   = new.get("Category", "")
        return f"Added '{title}' [{cat}]"

    if action in ("delete", "archive"):
        title = _trunc((old or new).get("Regulation_Title", ""), 50)
        return f"Archived '{title}'"

    if action == "edit":
        skip = {"Entry_ID", "entry_id", "HasData"}
        all_keys = (set(old) | set(new)) - skip
        ordered = _AUDIT_KEY_FIELDS + sorted(k for k in all_keys if k not in _AUDIT_KEY_FIELDS)
        parts: list[str] = []
        for k in ordered:
            o = str(old.get(k) or "").strip()
            n = str(new.get(k) or "").strip()
            if o != n:
                label = k.replace("_std", "").replace("_", " ")
                parts.append(f"{label}: {_trunc(o) or '—'} → {_trunc(n) or '—'}")
            if len(parts) == 5:
                break
        return "; ".join(parts) if parts else "No changes detected"

    return ""


def append_audit_log(
    action: str,
    country: str,
    entry_id: str,
    user: str,
    old_record: Optional[dict],
    new_record: Optional[dict],
) -> None:
    ARCHIVE_FILE.parent.mkdir(parents=True, exist_ok=True)
    row = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "action": action,
        "country": country,
        "entry_id": entry_id,
        "user": user,
        "old_record": (
            json.dumps(old_record, ensure_ascii=False)
            if old_record is not None
            else ""
        ),
        "new_record": (
            json.dumps(new_record, ensure_ascii=False)
            if new_record is not None
            else ""
        ),
    }
    header = not ARCHIVE_FILE.exists()
    pd.DataFrame([row]).to_csv(ARCHIVE_FILE, mode="a", index=False, header=header)


_REG_EXPORT_EXCLUDE = {"Entry_ID", "Year_raw", "Year_sort", "HasData"}
_REG_EXPORT_RENAME  = {
    "Country_std":      "Country",
    "Regulator_std":    "Regulator",
    "Regulation_Title": "Title",
    "Source_URL":       "Source URL",
    "Source_EN":        "Source URL (EN)",
    "Amendment_Of":     "Amendment Of",
}

def _build_regulations_excel(df: pd.DataFrame, by_category: bool = True) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        if by_category:
            for cat in sorted(df["Category"].dropna().unique()):
                sdf = (
                    df[df["Category"] == cat]
                    [[c for c in df.columns if c not in _REG_EXPORT_EXCLUDE]]
                    .rename(columns=_REG_EXPORT_RENAME)
                )
                sdf.to_excel(writer, sheet_name=str(cat)[:31], index=False)
        else:
            out = (
                df[[c for c in df.columns if c not in _REG_EXPORT_EXCLUDE]]
                .rename(columns=_REG_EXPORT_RENAME)
            )
            out.to_excel(writer, sheet_name="Regulations", index=False)
    return buf.getvalue()


def _docx_add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insert a clickable hyperlink run into an existing paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    # Inline formatting — don't rely on the "Hyperlink" named style which may not exist
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


def _build_country_word_doc(country: str, df: pd.DataFrame) -> bytes:
    """Generate a .docx report for a single country."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn

    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ────────────────────────────────────────────────────────────────
    flag = country_flag(country)
    doc.add_heading(f"{flag}  {country}", level=1)

    # ── Mini-map ─────────────────────────────────────────────────────────────
    try:
        _map_fig = _build_country_minimap(country)
        # Light background for the Word doc
        _map_fig.update_layout(
            paper_bgcolor="#FFFFFF",
            plot_bgcolor="#FFFFFF",
        )
        _map_fig.update_geos(
            landcolor="#D1D5DB",
            oceancolor="#EFF6FF",
            lakecolor="#EFF6FF",
            coastlinecolor="#9CA3AF",
            countrycolor="#6B7280",
        )
        _map_png = _map_fig.to_image(format="png", width=700, height=280, scale=2)
        doc.add_picture(io.BytesIO(_map_png), width=Inches(5.5))
        doc.paragraphs[-1].alignment = 0  # left-align the image paragraph
    except Exception:
        pass  # kaleido unavailable or figure error — skip map silently

    d = df[df["Country_std"] == country].copy()
    if d.empty:
        doc.add_paragraph("No regulations found for this country under the current filters.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Regulator ────────────────────────────────────────────────────────────
    regs = sorted(set(d["Regulator_std"].dropna().astype(str).tolist()))
    regs = [r for r in regs if r.lower() not in ("nan", "none", "")]
    p = doc.add_paragraph()
    p.add_run("Regulator: ").bold = True
    p.add_run(", ".join(regs) if regs else "—")

    # ── Regulations list ─────────────────────────────────────────────────────
    doc.add_heading("Regulations", level=2)
    latest = d.sort_values(["Year", "Regulation_Title"], ascending=[False, True])
    for _, row in latest.iterrows():
        year_text  = format_year(row.get("Year"))
        title_text = str(row.get("Regulation_Title") or "—").strip()
        source_url = safe_linkify(row.get("Source_URL"))
        source_en  = safe_linkify(row.get("Source_EN")) if pd.notna(row.get("Source_EN")) else ""
        amend      = row.get("Amendment_Of")
        amend_text = (
            f"  [amends: {amend}]"
            if pd.notna(amend) and str(amend).strip().lower() not in ("nan", "none", "")
            else ""
        )

        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{year_text}").bold = True
        p.add_run(f" — {title_text}{amend_text}")
        if source_url:
            p.add_run("  ")
            _docx_add_hyperlink(p, source_url, "[Source]")
        if source_en and source_en != source_url:
            p.add_run("  ")
            _docx_add_hyperlink(p, source_en, "[EN]")

    # ── Key provisions ────────────────────────────────────────────────────────
    doc.add_heading("Key Provisions", level=2)
    shown_any = False
    for cat in sorted(d["Category"].dropna().unique().tolist()):
        dc = d[d["Category"] == cat].copy()
        detail_cols = [
            c for c in dc.columns
            if c not in _KNOWN_META and c not in META_COL_CANDIDATES["source"]
        ]
        rows = []
        for col in detail_cols:
            vals = (
                dc[col].fillna(pd.NA).dropna().astype(str).str.strip()
            )
            unique_vals = sorted({
                v for v in vals
                if v.lower() not in ("nan", "none", "")
            })
            if unique_vals:
                rows.append((col, "; ".join(unique_vals)))
        if not rows:
            continue

        doc.add_heading(cat, level=3)
        tbl = doc.add_table(rows=len(rows) + 1, cols=2)
        tbl.style = "Table Grid"

        # Header row
        hdr = tbl.rows[0].cells
        for cell, label in zip(hdr, ("Field", "Values")):
            cell.text = label
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell._tc.get_or_add_tcPr().append(
                _docx_shading_elm("1E293B")
            )

        # Data rows
        for i, (field, value) in enumerate(rows):
            cells = tbl.rows[i + 1].cells
            cells[0].text = field
            cells[1].text = value
            if i % 2 == 0:
                for cell in cells:
                    cell._tc.get_or_add_tcPr().append(_docx_shading_elm("F1F5F9"))

        # Column widths: narrow field col, wide value col
        for row in tbl.rows:
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.3)

        shown_any = True

    if not shown_any:
        doc.add_paragraph("No key provisions available under the current filters.")

    doc.add_paragraph(
        "Source links point to the official regulation document. "
        "EN links open the English translation where available."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_shading_elm(fill_hex: str):
    """Return a w:shd OxmlElement for a solid table-cell background colour."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.lstrip("#"))
    return shd



def get_query_params() -> dict[str, str]:
    """Return current URL query parameters as a plain str->str dict.

    Streamlit 1.28+ exposes st.query_params as a Mapping[str, str].
    Older builds had experimental_get_query_params() -> dict[str, list[str]];
    we normalise both shapes to a flat str->str dict.
    """
    # Modern API (Streamlit >= 1.28)
    modern = getattr(st, "query_params", None)
    if modern is not None and not callable(modern):
        # st.query_params is a Mapping-like object, not a function
        return {k: str(v) for k, v in modern.items()}
    # Legacy API
    getter = getattr(st, "experimental_get_query_params", None)
    if callable(getter):
        raw: dict[str, list[str]] = getter()  # type: ignore[assignment]
        return {k: v[0] if v else "" for k, v in raw.items()}
    return {}


def save_cbregs(df: pd.DataFrame, src_dir: Path) -> None:
    src_dir = Path(src_dir)
    src_dir.mkdir(parents=True, exist_ok=True)
    for category, group in df.groupby("Category", sort=False):
        csv_path = src_dir / f"{category}.csv"
        group.to_csv(csv_path, index=False)


# =========================
# Gap Analysis data layer
# =========================
_GAP_BENCH_FILE   = "gap_benchmarks.csv"
_GAP_MAP_FILE     = "gap_mappings.csv"
_GAP_AUDIT_FILE   = "gap_audit_log.csv"

ASEAN_COUNTRIES_ORDERED = [
    "Brunei Darussalam", "Cambodia", "Indonesia", "Lao PDR",
    "Malaysia", "Myanmar", "Philippines", "Singapore",
    "Thailand", "Timor-Leste", "Viet Nam",
]

GAP_STATUS_OPTIONS = ["Meets", "Partially meets", "Does not meet", "Not assessed"]
GAP_STATUS_META: dict[str, dict] = {
    "Meets":           {"emoji": "✅", "color": _C_GREEN,  "bg": "#14532d", "text": "#86efac"},
    "Partially meets": {"emoji": "⚠️", "color": _C_AMBER,  "bg": "#78350f", "text": "#fcd34d"},
    "Does not meet":   {"emoji": "✗",  "color": _C_RED,    "bg": "#450a0a", "text": "#fca5a5"},
    "Not assessed":    {"emoji": "—",  "color": _C_GRAY,   "bg": "#0f172a", "text": _C_FAINT},
}


def _status_badge(status: str) -> str:
    m = GAP_STATUS_META.get(status, GAP_STATUS_META["Not assessed"])
    c = m["color"]
    e = m["emoji"]
    return (
        f"<span style='background:{c}22;border:1px solid {c}55;"
        f"border-radius:20px;padding:3px 12px;font-size:0.82rem;"
        f"font-weight:600;color:{c}'>{e} {status}</span>"
    )
_GAP_BENCH_COLS = [
    "Benchmark_ID", "Standard", "Category", "Topic", "Provision",
    "Created_By", "Created_At", "Updated_By", "Updated_At",
]
_GAP_MAP_COLS = [
    "Mapping_ID", "Benchmark_ID", "Country", "Status",
    "Entry_IDs", "Gap_Assessment", "Updated_By", "Updated_At",
]


def get_gap_cache_key(src_dir: Path) -> str:
    b = src_dir / _GAP_BENCH_FILE
    m = src_dir / _GAP_MAP_FILE
    return f"{b.stat().st_mtime if b.exists() else 0}_{m.stat().st_mtime if m.exists() else 0}"


@st.cache_data
def load_gap_benchmarks(src_dir: str, cache_key: str) -> pd.DataFrame:
    path = Path(src_dir) / _GAP_BENCH_FILE
    if not path.exists():
        return pd.DataFrame(columns=_GAP_BENCH_COLS)
    return pd.read_csv(path, dtype=str).fillna("")


@st.cache_data
def load_gap_mappings(src_dir: str, cache_key: str) -> pd.DataFrame:
    path = Path(src_dir) / _GAP_MAP_FILE
    if not path.exists():
        return pd.DataFrame(columns=_GAP_MAP_COLS)
    df = pd.read_csv(path, dtype=str).fillna("")
    if "Notes" in df.columns and "Gap_Assessment" not in df.columns:
        df = df.rename(columns={"Notes": "Gap_Assessment"})
    return df


def save_gap_benchmarks(df: pd.DataFrame, src_dir: Path) -> None:
    df.to_csv(src_dir / _GAP_BENCH_FILE, index=False)


def save_gap_mappings(df: pd.DataFrame, src_dir: Path) -> None:
    df.to_csv(src_dir / _GAP_MAP_FILE, index=False)


def append_gap_audit(
    src_dir: Path, action: str, record_type: str,
    record_id: str, before: dict, after: dict, user: str,
) -> None:
    path = src_dir / _GAP_AUDIT_FILE
    row = {
        "Timestamp":   pd.Timestamp.utcnow().isoformat(),
        "User":        user,
        "Action":      action,
        "Record_Type": record_type,
        "Record_ID":   record_id,
        "Before":      json.dumps(before, ensure_ascii=False),
        "After":       json.dumps(after,  ensure_ascii=False),
    }
    pd.DataFrame([row]).to_csv(path, mode="a", header=not path.exists(), index=False)


def _gap_reg_options(df_all: pd.DataFrame, country: str, category: str = "") -> dict[str, str]:
    """Return {Entry_ID: display_label} for a country's regulations, optionally filtered by category.
    Filtered by category so users only see regulations relevant to the benchmark's domain."""
    mask = df_all["Country_std"] == country
    if category and category not in ("", "nan", "All"):
        mask &= df_all["Category"] == category
    opts: dict[str, str] = {}
    for _, r in df_all[mask].dropna(subset=["Regulation_Title"]).iterrows():
        eid   = str(r["Entry_ID"])
        title = str(r.get("Regulation_Title", "Untitled"))
        yr    = r.get("Year", "")
        yr_s  = f" ({int(float(yr))})" if pd.notna(yr) and str(yr) not in ("", "nan") else ""
        opts[eid] = f"{title}{yr_s}"
    return opts


def _render_mapping_card(
    cur_status: str, cur_eids: list, cur_notes: str,
    is_new_map: bool, em_row, df_all: pd.DataFrame,
) -> None:
    """Render the read-only mapping detail card + optional edit button."""
    sc = GAP_STATUS_META.get(cur_status, GAP_STATUS_META["Not assessed"])["color"]
    if cur_eids:
        _linked = df_all[df_all["Entry_ID"].isin(cur_eids)]
        _pills = []
        for _, _rr in _linked.iterrows():
            _t   = str(_rr.get("Regulation_Title", "Regulation"))
            _url = str(_rr.get("Source_URL", ""))
            _yr  = str(_rr.get("Year", ""))
            _yr_s = (
                f" <span style='color:#94a3b8;font-size:0.78rem'>({int(float(_yr))})</span>"
                if _yr and _yr not in ("nan", "") else ""
            )
            _inner = (
                f'<a href="{_url}" target="_blank" style="color:#93c5fd;text-decoration:none">{_t}</a>{_yr_s}'
                if _url and _url not in ("", "nan") else f"{_t}{_yr_s}"
            )
            _pills.append(
                f"<div style='background:#1e293b;border:1px solid #334155;border-radius:6px;"
                f"padding:6px 10px;margin-bottom:5px;font-size:0.85rem;color:#e2e8f0;"
                f"line-height:1.4'>{_inner}</div>"
            )
        regs_html = "".join(_pills)
    else:
        regs_html = "<span style='color:#64748b;font-size:0.83rem;font-style:italic'>No regulations linked yet.</span>"

    notes_html = ""
    if cur_notes or cur_status in ("Partially meets", "Does not meet"):
        _nb = (
            cur_notes if cur_notes
            else "<em style='color:#64748b'>No gap assessment provided — edit the mapping to describe what is still needed.</em>"
        )
        notes_html = (
            "<div style='margin-top:14px'>"
            "<div style='font-size:0.72rem;font-weight:600;letter-spacing:0.07em;"
            "text-transform:uppercase;color:#94a3b8;margin-bottom:5px'>Gap assessment</div>"
            "<div style='background:#1e293b;border-left:3px solid #f59e0b;"
            f"border-radius:0 6px 6px 0;padding:10px 12px;font-size:0.85rem;color:#e2e8f0;line-height:1.5'>{_nb}</div>"
            "</div>"
        )

    upd_html = ""
    if not is_new_map and em_row is not None:
        upd_html = (
            f"<div style='font-size:0.72rem;color:#475569;margin-top:12px'>"
            f"Last updated {str(em_row.get('Updated_At',''))[:10]} by {em_row.get('Updated_By','')}</div>"
        )

    st.markdown(
        f"<div style='background:#0f172a;border:1px solid #1e293b;border-radius:10px;"
        f"padding:16px 18px;margin:10px 0'>"
        f"<div style='display:flex;align-items:center;gap:10px;margin-bottom:14px'>"
        f"<span style='background:{sc}22;border:1px solid {sc}55;border-radius:20px;padding:3px 12px;"
        f"font-size:0.82rem;font-weight:600;color:{sc}'>"
        f"{GAP_STATUS_META.get(cur_status, GAP_STATUS_META['Not assessed'])['emoji']} {cur_status}</span></div>"
        f"<div style='font-size:0.72rem;font-weight:600;letter-spacing:0.07em;"
        f"text-transform:uppercase;color:#94a3b8;margin-bottom:7px'>Linked regulations</div>"
        f"{regs_html}{notes_html}{upd_html}</div>",
        unsafe_allow_html=True,
    )

    # edit button rendered by caller (needs dialog call with full context)


def _render_gap_matrix(df_bench: pd.DataFrame, df_map: pd.DataFrame, countries: list[str]) -> None:
    """Render the gap analysis status matrix: benchmarks (rows) × countries (cols)."""
    import html as _hl
    if df_bench.empty:
        st.info("No benchmarks have been defined yet. Admin can add them in the Benchmarks tab.")
        return

    status_lk: dict[tuple[str, str], str] = {}
    if not df_map.empty:
        for _, r in df_map.iterrows():
            status_lk[(str(r["Benchmark_ID"]), str(r["Country"]))] = str(r.get("Status", "Not assessed"))

    short_name = {
        "Brunei Darussalam": "Brunei", "Timor-Leste": "Timor-Leste",
        "Lao PDR": "Lao PDR",
    }

    th = (
        '<th style="min-width:130px">Standard</th>'
        '<th style="min-width:180px">Topic / Area</th>'
    )
    for c in countries:
        flag = country_flag(c)
        label = _hl.escape(short_name.get(c, c))
        th += (
            f'<th style="text-align:center;min-width:70px;font-size:11px;line-height:1.4">'
            f'{flag}<br>{label}</th>'
        )

    rows_html = ""
    for i, (_, b) in enumerate(df_bench.iterrows()):
        bid   = str(b.get("Benchmark_ID", ""))
        std   = str(b.get("Standard", ""))
        topic = str(b.get("Topic", ""))
        bg    = "#0d1526" if i % 2 == 0 else "#111e33"
        cells = (
            f'<td style="font-size:11px;color:#94a3b8">{_hl.escape(std)}</td>'
            f'<td style="font-size:13px">{_hl.escape(topic)}</td>'
        )
        for c in countries:
            status = status_lk.get((bid, c), "Not assessed")
            _sm    = GAP_STATUS_META.get(status, GAP_STATUS_META["Not assessed"])
            cells += (
                f'<td style="background:{_sm["bg"]};color:{_sm["text"]};'
                f'text-align:center;font-size:16px;padding:5px 3px">{_sm["emoji"]}</td>'
            )
        rows_html += f'<tr style="background:{bg}">{cells}</tr>'

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:transparent;font-family:"Source Sans Pro",sans-serif;color:#e2e8f0}}
table{{width:100%;border-collapse:collapse}}
thead tr{{background:#1e293b;position:sticky;top:0;z-index:2}}
th{{padding:7px 10px 7px 8px;text-align:left;color:#94a3b8;font-size:13px;font-weight:600;border-bottom:2px solid #334155;vertical-align:middle}}
td{{padding:6px 10px;border-bottom:1px solid #1e293b;vertical-align:middle;word-break:break-word;font-size:13.5px}}
tr:hover td{{filter:brightness(1.2)}}
@media print{{body,table,thead tr,td,th{{background:#fff!important;color:#111!important;border-color:#ccc!important;filter:none!important}}}}
</style></head><body>
<table id="t"><thead><tr>{th}</tr></thead><tbody>{rows_html}</tbody></table>
<script>
(function(){{const tbl=document.getElementById('t');let _lH=0;
function sh(){{const h=Math.ceil(tbl.getBoundingClientRect().height)+8;
if(h>0&&h!==_lH){{_lH=h;window.parent.postMessage({{isStreamlitMessage:true,type:'streamlit:setFrameHeight',height:h}},'*');}}}}
sh();window.addEventListener('load',sh);setTimeout(sh,100);setTimeout(sh,500);
new ResizeObserver(sh).observe(tbl);}})();
</script></body></html>"""

    # 46px header + 42px per row + 16px wrapper — no large floor to avoid blank gap
    est = max(80, 62 + len(df_bench) * 42)
    st.iframe(src=_html_src(html), height=est)


@st.dialog("Edit regulation", width="large")
def edit_regulation_dialog(
    existing_row: pd.Series,
    category_choices: List[str],
    df_all: pd.DataFrame,
    src_dir: Path,
    editor_country: str,
    auth_user: str,
) -> None:
    selected_entry_id = str(existing_row["Entry_ID"])
    save_key = f"edit_saved_{selected_entry_id}"
    confirm_key = f"confirm_delete_{selected_entry_id}"
    st.session_state.setdefault(save_key, False)
    st.session_state.setdefault(confirm_key, False)

    existing_category = (
        existing_row["Category"]
        if pd.notna(existing_row["Category"])
        else category_choices[0]
    )
    edit_provision_cols = get_provision_cols(df_all, existing_category)

    def _s(val) -> str:
        return "" if pd.isna(val) else str(val)

    # Reset field keys when dialog opens for a different entry
    if st.session_state.get("_dialog_entry_id") != selected_entry_id:
        st.session_state["_dialog_entry_id"] = selected_entry_id
        for k in ["_dlg_cat", "_dlg_reg", "_dlg_year", "_dlg_title", "_dlg_url", "_dlg_url_en", "_dlg_amend"]:
            st.session_state.pop(k, None)
        for col in edit_provision_cols:
            st.session_state.pop(f"_dlg_prov_{col}", None)

    # Editable fields (no form — required for real-time change detection)
    edit_category = st.selectbox(
        "Category",
        options=category_choices,
        index=category_choices.index(existing_category) if existing_category in category_choices else 0,
        key="_dlg_cat",
    )
    edit_regulator  = st.text_input("Regulator",           value=_s(existing_row.get("Regulator_std")),    key="_dlg_reg")
    edit_year_raw   = st.text_input("Year",                value=_s(existing_row.get("Year_raw")),         key="_dlg_year")
    edit_title      = st.text_area("Regulation title",     value=_s(existing_row.get("Regulation_Title")), height=120, key="_dlg_title")
    edit_source_url = st.text_input("Source URL",          value=_s(existing_row.get("Source_URL")),       key="_dlg_url")
    edit_source_en  = st.text_input("Source URL (English)", value=_s(existing_row.get("Source_EN")),       key="_dlg_url_en")

    _amend_country = str(existing_row.get("Country_std", ""))
    _amend_cand = df_all[
        (df_all["Country_std"] == _amend_country) &
        (df_all["Category"] == edit_category) &
        df_all["Regulation_Title"].notna() &
        (df_all["Entry_ID"] != selected_entry_id)
    ].sort_values(["Year", "Regulation_Title"], ascending=[False, True])
    _BLANK_OPT = "— (not an amendment)"
    amend_opts = [_BLANK_OPT] + [
        f"{r['Regulation_Title']} ({format_year(r['Year'])})"
        for _, r in _amend_cand.iterrows()
    ]
    _cur_amend = _s(existing_row.get("Amendment_Of"))
    _amend_idx = amend_opts.index(_cur_amend) if _cur_amend in amend_opts else 0
    edit_amendment_of_sel = st.selectbox(
        "Amends (leave blank if original)",
        options=amend_opts,
        index=_amend_idx,
        key="_dlg_amend",
    )
    edit_amendment_of = "" if edit_amendment_of_sel == _BLANK_OPT else edit_amendment_of_sel

    edit_provision_values: dict[str, str] = {}
    if edit_provision_cols:
        st.markdown("**Key provisions**")
        for col in edit_provision_cols:
            edit_provision_values[col] = st.text_area(
                col, value=_s(existing_row.get(col)), height=68, key=f"_dlg_prov_{col}"
            )

    # Detect changes against the original row
    orig = {
        "_dlg_cat":    existing_category,
        "_dlg_reg":    _s(existing_row.get("Regulator_std")),
        "_dlg_year":   _s(existing_row.get("Year_raw")),
        "_dlg_title":  _s(existing_row.get("Regulation_Title")),
        "_dlg_url":    _s(existing_row.get("Source_URL")),
        "_dlg_url_en": _s(existing_row.get("Source_EN")),
        "_dlg_amend":  _cur_amend if _cur_amend in amend_opts else _BLANK_OPT,
        **{f"_dlg_prov_{col}": _s(existing_row.get(col)) for col in edit_provision_cols},
    }
    has_changes = any(st.session_state.get(k, orig[k]) != orig[k] for k in orig)

    if st.button("Save updates", type="primary", disabled=not has_changes, width="stretch"):
        try:
            updated_df = df_all.copy()
            row_index = updated_df.index[updated_df["Entry_ID"] == selected_entry_id][0]
            old_record = serialize_record_for_archive(updated_df.loc[row_index])
            updated_df.at[row_index, "Category"]      = edit_category
            updated_df.at[row_index, "Regulator_std"] = edit_regulator.strip() or pd.NA
            updated_df.at[row_index, "Year_raw"]      = edit_year_raw.strip() or pd.NA
            updated_df.at[row_index, "Year"]          = extract_year(edit_year_raw)
            updated_df.at[row_index, "Regulation_Title"] = edit_title.strip()
            updated_df.at[row_index, "Source_URL"]    = edit_source_url.strip() or pd.NA
            updated_df.at[row_index, "Source_EN"]     = edit_source_en.strip() or pd.NA
            updated_df.at[row_index, "Amendment_Of"]  = edit_amendment_of.strip() or pd.NA
            for col, val in edit_provision_values.items():
                updated_df.at[row_index, col] = val.strip() or pd.NA
            save_cbregs(updated_df, src_dir)
            append_audit_log(
                action="edit",
                country=editor_country,
                entry_id=selected_entry_id,
                user=auth_user,
                old_record=old_record,
                new_record=serialize_record_for_archive(updated_df.loc[row_index]),
            )
            st.session_state[save_key] = True
        except Exception as _e:
            st.error(f"Save failed: {_e}")

    if st.session_state[save_key]:
        st.success("Changes saved.")

    if st.button("Close", type="primary", width="stretch"):
        st.session_state.pop(save_key, None)
        for _cat in df_all["Category"].dropna().unique():
            _cat_key = re.sub(r"[^a-zA-Z0-9]", "_", str(_cat))
            st.session_state.pop(f"editor_table_{_cat_key}", None)
        st.rerun()

    # Marker + scoped CSS: only secondary buttons after this point in the dialog are red
    _arch_marker = f"dlg-arch-{selected_entry_id[:8]}"
    st.markdown(
        f'<div id="{_arch_marker}"></div>'
        f'<style>'
        f'div:has(>div#{_arch_marker}) ~ div [data-testid="baseButton-secondary"]'
        f'{{background-color:#dc2626!important;border-color:#dc2626!important;color:#fff!important}}'
        f'div:has(>div#{_arch_marker}) ~ div [data-testid="baseButton-secondary"]:hover'
        f'{{background-color:#b91c1c!important;border-color:#b91c1c!important}}'
        f'</style>',
        unsafe_allow_html=True,
    )
    st.divider()
    if not st.session_state[confirm_key]:
        if st.button("Archive and delete this record", type="secondary", width="stretch"):
            st.session_state[confirm_key] = True
            st.rerun()
    else:
        st.warning("Are you sure? This record will be permanently removed from the dataset.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes, delete", type="secondary", width="stretch"):
                updated_df = df_all.copy()
                row_index = updated_df.index[updated_df["Entry_ID"] == selected_entry_id][0]
                old_record = serialize_record_for_archive(updated_df.loc[row_index])
                updated_df = updated_df.drop(index=row_index).reset_index(drop=True)
                save_cbregs(updated_df, src_dir)
                append_audit_log(
                    action="delete",
                    country=editor_country,
                    entry_id=selected_entry_id,
                    user=auth_user,
                    old_record=old_record,
                    new_record=None,
                )
                st.session_state.pop(confirm_key, None)
                st.rerun()
        with col2:
            if st.button("Cancel", type="primary", width="stretch"):
                st.session_state[confirm_key] = False
                st.rerun()


# =========================
# Load data
# =========================
resolved_data_dir = resolve_data_dir(DATA_DIR)

df_all = load_cbregs(str(resolved_data_dir), get_csv_cache_key(resolved_data_dir))

# Persist Entry_IDs to CSVs if any are missing (ensures gap mapping references stay stable).
# Only runs once per session — guarded so it doesn't repeat on every rerun.
if not st.session_state.get("_entry_ids_saved"):
    _needs_save = any(
        "Entry_ID" not in pd.read_csv(f, nrows=0).columns
        for f in sorted(resolved_data_dir.glob("*.csv"))
    )
    if _needs_save:
        save_cbregs(df_all, resolved_data_dir)
        st.cache_data.clear()
        st.rerun()
    st.session_state["_entry_ids_saved"] = True

if df_all.empty:
    st.error("Data loaded but produced no rows.")
    st.stop()

# =========================
# Category filter + download (above tabs)
# =========================
categories = ["All"] + sorted(df_all["Category"].dropna().unique().tolist())

_CAT_KEY_MAP   = "cat_filter_map"
_CAT_KEY_TABLE = "cat_filter_table"


def _on_map_category_change():
    st.session_state[_CAT_KEY_TABLE] = st.session_state.get(_CAT_KEY_MAP) or "All"
    st.session_state.pop("table_a_country_sel", None)
    st.session_state.pop("table_b_country_sel", None)
    st.session_state.pop("map_country_sel", None)


def _on_table_category_change():
    st.session_state[_CAT_KEY_MAP] = st.session_state.get(_CAT_KEY_TABLE) or "All"
    st.session_state.pop("table_a_country_sel", None)
    st.session_state.pop("table_b_country_sel", None)
    st.session_state.pop("map_country_sel", None)


def _render_category_filter(tab: str) -> str:
    """Render category dropdown + export button inline; syncs state between Map and Table."""
    key       = _CAT_KEY_MAP if tab == "map" else _CAT_KEY_TABLE
    on_change = _on_map_category_change if tab == "map" else _on_table_category_change
    cur_val   = st.session_state.get(key) or "All"
    cur_idx   = categories.index(cur_val) if cur_val in categories else 0
    _col_sel, _, _col_dl = st.columns([1, 4, 1], vertical_alignment="bottom")
    _sel = _col_sel.selectbox(
        "Category",
        options=categories,
        index=cur_idx,
        key=key,
        on_change=on_change,
    )
    # Build export: apply current category + any selected countries from session state
    _exp = df_all[df_all["Category"] == _sel].copy() if _sel != "All" else df_all.copy()
    if tab == "table":
        _country_key = "table_a_country_sel" if _sel == "All" else "table_b_country_sel"
        _sel_countries = st.session_state.get(_country_key) or []
        if _sel_countries:
            _exp = _exp[_exp["Country_std"].isin(_sel_countries)]
    _exp_clean = (
        _exp[[c for c in _exp.columns if c not in _REG_EXPORT_EXCLUDE]]
        .rename(columns=_REG_EXPORT_RENAME)
    )
    _col_dl.download_button(
        "⬇ Export", _exp_clean.to_csv(index=False).encode("utf-8"),
        "asean_regulations_filtered.csv", "text/csv",
        width="stretch", key=f"dl_export_{tab}",
    )
    return _sel


# =========================
# Country extender (modal)
# =========================
def _build_country_minimap(country: str) -> "go.Figure":
    """Return a Plotly choropleth figure highlighting a single ASEAN country."""
    map_df = pd.DataFrame({
        "ISO3":     [ASEAN_ISO3.get(country, "")],
        "Country":  [country],
        "Selected": [1],
    })
    fig = px.choropleth(
        map_df,
        locations="ISO3",
        locationmode="ISO-3",
        color="Selected",
        color_continuous_scale=["#facc15", "#facc15"],
        range_color=(0, 1),
        hover_name="Country",
    )
    fig.update_traces(showscale=False, hovertemplate="<extra></extra>")
    fig.update_geos(
        scope="asia",
        projection_type="mercator",
        lonaxis=dict(range=[92, 141]),
        lataxis=dict(range=[-11, 24]),
        visible=True,
        showcoastlines=True,
        coastlinecolor="rgba(255,255,255,0.35)",
        showcountries=True,
        countrycolor="rgba(255,255,255,0.85)",
        showland=True,
        landcolor="rgba(20, 30, 45, 1)",
        showocean=True,
        oceancolor="rgba(10, 16, 26, 1)",
        showlakes=True,
        lakecolor="rgba(10, 16, 26, 1)",
        bgcolor="rgba(0,0,0,0)",
    )
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=0, b=0),
        height=220,
        showlegend=False,
        coloraxis_showscale=False,
    )
    return fig


def show_country_modal(country: str, df: pd.DataFrame, key_suffix: str = "", show_minimap: bool = True, show_pdf_btn: bool = False):
    st.divider()
    st.markdown(f"### {country_flag(country)} {country}")

    if show_minimap:
        st.plotly_chart(
            _build_country_minimap(country),
            width="stretch",
            key=f"country_modal_chart_{country}_{key_suffix}",
            config={
                "displayModeBar": False,
                "scrollZoom": False,
                "doubleClick": False,
                "responsive": True,
            },
        )

    d = df[df["Country_std"] == country].copy()
    if d.empty:
        st.info(
            "No regulations found for this country under the current filters."
        )
        return

    # Download buttons
    _dl_multi_cat = d["Category"].dropna().nunique() > 1 if "Category" in d.columns else False
    _xl_bytes = _build_regulations_excel(d, by_category=_dl_multi_cat)
    _xl_name  = f"{country.lower().replace(' ', '_')}_regulations.xlsx"
    _dl_col1, _dl_col2 = st.columns(2) if show_pdf_btn else (st, None)
    _dl_col1.download_button(
        label=f"⬇ Export (.xlsx)",
        data=_xl_bytes,
        file_name=_xl_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        width="stretch",
        key=f"dl_{country}_{key_suffix}",
    )
    if show_pdf_btn and _dl_col2 is not None:
        _word_bytes = _build_country_word_doc(country, df)
        _word_name  = f"{country.lower().replace(' ', '_')}_report.docx"
        _dl_col2.download_button(
            label="📄 Download as Word",
            data=_word_bytes,
            file_name=_word_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width="stretch",
            key=f"word_{country}_{key_suffix}",
        )

    regs = sorted(set(d["Regulator_std"].dropna().tolist()))
    if regs:
        st.markdown("**Regulator:** " + ", ".join(regs))
    else:
        st.markdown("**Regulator:** —")

    st.markdown("**Regulations**")
    latest = d.sort_values(["Year", "Regulation_Title"], ascending=[False, True])
    regulation_lines = []
    for _, row in latest.iterrows():
        year = row.get("Year")
        year_text = format_year(year)
        title = str(row.get("Regulation_Title", "—"))
        source_link = safe_linkify(row.get("Source_URL"))
        source_en_raw = row.get("Source_EN")
        source_en_link = safe_linkify(source_en_raw) if pd.notna(source_en_raw) else ""
        amendment_of = row.get("Amendment_Of")

        src_parts = []
        if source_link:
            src_parts.append(f"[Source]({source_link})")
        if source_en_link and source_en_link != source_link:
            src_parts.append(f"[EN]({source_en_link})")
        source_text = f" ({' · '.join(src_parts)})" if src_parts else ""

        amend_text = ""
        if pd.notna(amendment_of) and str(amendment_of).strip().lower() not in ("nan", "none", ""):
            amend_text = f" *(amends {amendment_of})*"

        regulation_lines.append(f"- **{year_text}** — {title}{source_text}{amend_text}")
    st.markdown("\n".join(regulation_lines))

    st.markdown("**Key provisions**")
    known_meta_cols = _KNOWN_META

    shown_any_category = False
    for cat in sorted(d["Category"].dropna().unique().tolist()):
        dc = d[d["Category"] == cat].copy()
        detail_cols = [
            c for c in dc.columns
            if c not in known_meta_cols
            and c not in META_COL_CANDIDATES["source"]
        ]
        if not detail_cols:
            continue

        rows = []
        for col in detail_cols:
            values = (
                dc[col]
                .fillna(pd.NA)
                .dropna()
                .astype(str)
                .str.strip()
            )
            unique_values = sorted({
                v for v in values
                if pd.notna(v) and v.lower() not in {"nan", "none", ""}
            })
            if not unique_values:
                continue
            rows.append({"Field": col, "Values": "; ".join(unique_values)})

        if not rows:
            continue

        st.markdown(f"#### {cat}")
        _render_provision_table(rows, ["Field", "Values"])
        shown_any_category = True

    if not shown_any_category:
        st.caption(
            "No key provisions available for this country "
            "under the current filters."
        )

    st.caption(
        "Source links point to the official regulation document. "
        "EN links open the English translation where available."
    )


def show_country_comparison(countries: List[str], df: pd.DataFrame, key_suffix: str = ""):
    """Side-by-side Key provisions comparison for two or more jurisdictions."""
    st.divider()
    flags_line = "  ·  ".join(
        f"{country_flag(c)} {c}" for c in countries
    )
    st.markdown(f"## Comparing: {flags_line}")

    known_meta_cols = _KNOWN_META

    # Column header = "flag ISO" e.g. "🇸🇬 SG"
    def col_label(c: str) -> str:
        iso = COUNTRY_ISO_CODES.get(c, c[:2].upper())
        return f"{country_flag(c)} {iso}"

    col_labels = [col_label(c) for c in countries]

    # Gather all categories that appear for any of the selected countries
    all_cats = sorted(
        df[df["Country_std"].isin(countries)]["Category"].dropna().unique().tolist()
    )

    shown_any_category = False
    for cat in all_cats:
        # Collect detail columns across all countries for this category
        detail_cols_set: list[str] = []
        country_data: dict[str, pd.DataFrame] = {}
        for c in countries:
            dc = df[(df["Country_std"] == c) & (df["Category"] == cat)].copy()
            country_data[c] = dc
            for col in dc.columns:
                if (
                    col not in known_meta_cols
                    and col not in META_COL_CANDIDATES["source"]
                    and col not in detail_cols_set
                ):
                    detail_cols_set.append(col)

        if not detail_cols_set:
            continue

        # Build comparison rows: Field | <flag ISO col per country>
        rows = []
        for col in detail_cols_set:
            row: dict[str, str] = {"Field": col}
            has_any_value = False
            for c, label in zip(countries, col_labels):
                dc = country_data[c]
                if col not in dc.columns:
                    row[label] = "—"
                    continue
                values = (
                    dc[col]
                    .fillna(pd.NA)
                    .dropna()
                    .astype(str)
                    .str.strip()
                )
                unique_values = sorted({
                    v for v in values
                    if pd.notna(v) and v.lower() not in {"nan", "none", ""}
                })
                if unique_values:
                    row[label] = "; ".join(unique_values)
                    has_any_value = True
                else:
                    row[label] = "—"
            if has_any_value:
                rows.append(row)

        if not rows:
            continue

        st.markdown(f"#### {cat}")
        _render_provision_table(rows, ["Field"] + col_labels)
        shown_any_category = True

    if not shown_any_category:
        st.caption(
            "No key provisions found for the selected jurisdictions "
            "under the current filters."
        )

    st.caption(
        "Column headers show the flag and ISO code for each jurisdiction."
    )


@st.dialog("Add regulation", width="large")
def add_regulation_dialog(
    category_choices: List[str],
    df_all: pd.DataFrame,
    src_dir: Path,
    editor_country: str,
    auth_user: str,
) -> None:
    existing_regulators = sorted(
        set(
            df_all[df_all["Country_std"] == editor_country]["Regulator_std"]
            .dropna().astype(str).str.strip().tolist()
        )
    )
    prefilled_regulator = "; ".join(existing_regulators)

    _CAT_PLACEHOLDER = "— Select a category —"
    new_category = st.selectbox(
        "Category",
        options=[_CAT_PLACEHOLDER] + category_choices,
        index=0,
        key="_add_dlg_cat",
    )
    category_chosen = new_category != _CAT_PLACEHOLDER

    add_provision_cols = get_provision_cols(df_all, new_category) if category_chosen else []

    _ADD_BLANK = "— (not an amendment)"
    if category_chosen:
        _add_amend_mask = (df_all["Category"] == new_category) & df_all["Regulation_Title"].notna()
        if editor_country and editor_country != "NA":
            _add_amend_mask &= df_all["Country_std"] == editor_country
        _add_amend_cand = df_all[_add_amend_mask].sort_values(
            ["Year", "Regulation_Title"], ascending=[False, True]
        )
        add_amend_opts = [_ADD_BLANK] + [
            f"{r['Regulation_Title']} ({int(r['Year']) if pd.notna(r['Year']) else '?'})"
            for _, r in _add_amend_cand.iterrows()
        ]
    else:
        add_amend_opts = [_ADD_BLANK]

    new_amendment_of_sel = st.selectbox(
        "Amends (leave blank if original)", options=add_amend_opts, key="_add_dlg_amend"
    )
    new_amendment_of = "" if new_amendment_of_sel == _ADD_BLANK else new_amendment_of_sel

    new_regulator   = st.text_input("Regulator", value=prefilled_regulator, disabled=True)
    new_year_raw    = st.text_input("Year", value="", key="_add_dlg_year")
    new_title       = st.text_area("Regulation title", value="", height=120, key="_add_dlg_title")
    new_source_url  = st.text_input("Source URL", value="", key="_add_dlg_url")
    new_source_en   = st.text_input("Source URL (English)", value="", key="_add_dlg_url_en")

    add_provision_values: dict[str, str] = {}
    if add_provision_cols:
        st.markdown("**Key provisions**")
        for col in add_provision_cols:
            add_provision_values[col] = st.text_area(
                col, value="", height=68, key=f"_add_dlg_prov_{col}"
            )

    if st.button("Add regulation", type="primary", width="stretch"):
        if not category_chosen:
            st.warning("Please select a category.")
        elif not new_title.strip():
            st.warning("A regulation title is required.")
        else:
            new_row: dict[str, object] = {col: pd.NA for col in df_all.columns}
            new_row["Entry_ID"]        = str(uuid.uuid4())
            new_row["Category"]        = new_category
            new_row["Country_std"]     = editor_country
            new_row["Regulator_std"]   = new_regulator.strip() or pd.NA
            new_row["Year_raw"]        = new_year_raw.strip() or pd.NA
            new_row["Year"]            = extract_year(new_year_raw)
            new_row["Regulation_Title"] = new_title.strip()
            new_row["Source_URL"]      = new_source_url.strip() or pd.NA
            new_row["Source_EN"]       = new_source_en.strip() or pd.NA
            new_row["Amendment_Of"]    = new_amendment_of.strip() or pd.NA
            for col, val in add_provision_values.items():
                new_row[col] = val.strip() or pd.NA
            try:
                updated_df = pd.concat([df_all, pd.DataFrame([new_row])], ignore_index=True, sort=False)
                save_cbregs(updated_df, src_dir)
                append_audit_log(
                    action="add",
                    country=editor_country,
                    entry_id=str(new_row["Entry_ID"]),
                    user=auth_user,
                    old_record=None,
                    new_record=serialize_record_for_archive(pd.Series(new_row)),
                )
                st.success("Regulation saved.")
                st.rerun()
            except Exception as _e:
                st.error(f"Save failed: {_e}")

# =========================
# =========================
# Gap Analysis dialogs
# =========================
@st.dialog("Add benchmark", width="large")
def _dlg_add_benchmark(df_bench: pd.DataFrame, categories: list[str], src_dir: Path, user: str) -> None:
    bid   = st.text_input("Benchmark ID *", placeholder="e.g. AML-01")
    std   = st.text_input("Standard", placeholder="e.g. FATF Recommendations")
    cat   = st.selectbox("Regulatory category", options=[""] + categories)
    topic = st.text_input("Topic / Area")
    prov  = st.text_area("Provision text", height=220)
    if st.button("Add benchmark", type="primary"):
        if not bid.strip():
            st.error("Benchmark ID is required.")
            return
        if bid.strip() in df_bench["Benchmark_ID"].tolist():
            st.error(f"Benchmark ID '{bid.strip()}' already exists.")
            return
        now = pd.Timestamp.utcnow().isoformat()
        new_row = {
            "Benchmark_ID": bid.strip(), "Standard": std.strip(), "Category": cat,
            "Topic": topic.strip(), "Provision": prov.strip(),
            "Created_By": user, "Created_At": now, "Updated_By": user, "Updated_At": now,
        }
        try:
            save_gap_benchmarks(pd.concat([df_bench, pd.DataFrame([new_row])], ignore_index=True), src_dir)
            append_gap_audit(src_dir, "add", "benchmark", bid.strip(), {}, new_row, user)
            st.rerun()
        except Exception as _e:
            st.error(f"Save failed: {_e}")


@st.dialog("Edit benchmark", width="large")
def _dlg_edit_benchmark(row: pd.Series, df_bench: pd.DataFrame, categories: list[str], src_dir: Path, user: str) -> None:
    bid   = str(row["Benchmark_ID"])
    st.markdown(f"**Benchmark ID:** {bid}")
    std   = st.text_input("Standard",     value=str(row.get("Standard", "")))
    cur_cat = str(row.get("Category", ""))
    cat_opts = [""] + categories
    cat   = st.selectbox("Regulatory category", options=cat_opts,
                         index=cat_opts.index(cur_cat) if cur_cat in cat_opts else 0)
    topic = st.text_input("Topic / Area", value=str(row.get("Topic", "")))
    prov  = st.text_area("Provision text", value=str(row.get("Provision", "")), height=220)
    if st.button("Save", type="primary"):
        before = row.to_dict()
        now    = pd.Timestamp.utcnow().isoformat()
        idx    = df_bench.index[df_bench["Benchmark_ID"] == bid].tolist()
        if idx:
            for col, val in [("Standard", std.strip()), ("Category", cat),
                              ("Topic", topic.strip()), ("Provision", prov.strip()),
                              ("Updated_By", user), ("Updated_At", now)]:
                df_bench.at[idx[0], col] = val
            try:
                save_gap_benchmarks(df_bench, src_dir)
                append_gap_audit(src_dir, "edit", "benchmark", bid, before, df_bench.loc[idx[0]].to_dict(), user)
                st.rerun()
            except Exception as _e:
                st.error(f"Save failed: {_e}")
        else:
            st.rerun()


@st.dialog("Benchmark details", width="large")
def _dlg_benchmark_detail(
    brow: pd.Series,
    df_maps: pd.DataFrame,
    df_all: pd.DataFrame,
) -> None:
    bid = str(brow["Benchmark_ID"])

    _bench_maps = df_maps[df_maps["Benchmark_ID"] == bid]
    detail_rows:  list[dict] = []
    export_rows:  list[dict] = []
    for _c in ASEAN_COUNTRIES_ORDERED:
        _cm = _bench_maps[_bench_maps["Country"] == _c]
        if _cm.empty:
            continue
        _m = _cm.iloc[0]
        _status = str(_m.get("Status", "Not assessed"))
        if _status == "Not assessed":
            continue
        _eids = [e.strip() for e in str(_m.get("Entry_IDs", "")).split(",") if e.strip()]
        if not _eids:
            continue
        _linked = df_all[df_all["Entry_ID"].isin(_eids)]
        reg_lines, reg_texts = [], []
        for _, _rr in _linked.iterrows():
            _t   = str(_rr.get("Regulation_Title", "Regulation"))
            _url = str(_rr.get("Source_URL", ""))
            _yr  = str(_rr.get("Year", ""))
            _yr_s = f" ({int(float(_yr))})" if _yr and _yr not in ("nan", "") else ""
            reg_lines.append(
                f"[{_t}{_yr_s}]({_url})" if _url and _url not in ("", "nan") else f"{_t}{_yr_s}"
            )
            reg_texts.append(f"{_t}{_yr_s}")
        _explanation = str(_m.get("Gap_Assessment", _m.get("Notes", ""))).strip()
        if _explanation in ("", "nan"):
            _explanation = ""
        detail_rows.append({
            "Country":        f"{country_flag(_c)} {_c}",
            "Status":         GAP_STATUS_META.get(_status, GAP_STATUS_META["Not assessed"])["emoji"] + " " + _status,
            "Regulations":    "\n".join(f"• {r}" for r in reg_lines),
            "Gap assessment": _explanation,
        })
        export_rows.append({
            "Country":        _c,
            "Status":         _status,
            "Regulations":    "; ".join(reg_texts),
            "Gap assessment": _explanation,
        })

    # Build Excel for download
    _xl_buf = io.BytesIO()
    with pd.ExcelWriter(_xl_buf, engine="openpyxl") as _xl_w:
        pd.DataFrame([{
            "Benchmark ID": bid,
            "Standard":     str(brow.get("Standard", "")),
            "Category":     str(brow.get("Category", "")),
            "Topic":        str(brow.get("Topic", "")),
            "Provision":    str(brow.get("Provision", "")),
        }]).to_excel(_xl_w, sheet_name="Benchmark Details", index=False)
        pd.DataFrame(export_rows if export_rows else [
            {"Country": "", "Status": "", "Regulations": "", "Gap assessment": ""}
        ]).to_excel(_xl_w, sheet_name="Mapped Regulations", index=False)

    # Header row: title + download button
    _hdr_l, _hdr_r = st.columns([4, 1])
    _hdr_l.markdown(f"### {bid} — {brow.get('Topic', '')}")
    _hdr_r.download_button(
        "⬇ Export", _xl_buf.getvalue(),
        f"benchmark_{bid}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        width="stretch", key=f"dlg_bench_dl_{bid}",
    )
    st.caption(
        f"Standard: **{brow.get('Standard', '')}**  ·  "
        f"Category: **{brow.get('Category', '')}**"
    )
    with st.expander("Provision text", expanded=False):
        st.markdown(str(brow.get("Provision", "")))
    st.divider()

    if detail_rows:
        st.caption("Only jurisdictions with mapped regulations are shown.")
        _render_provision_table(detail_rows, ["Country", "Status", "Regulations", "Gap assessment"])
    else:
        st.info("No jurisdictions have tagged regulations for this benchmark yet.")


@st.dialog("Edit mapping", width="large")
def _dlg_edit_mapping(
    bid: str,
    map_country: str,
    mbrow: pd.Series,
    cur_status: str,
    cur_eids: list,
    cur_notes: str,
    mapping_id: str,
    is_new_map: bool,
    df_all: pd.DataFrame,
    df_maps: pd.DataFrame,
    gap_dir: Path,
    auth_user: str,
) -> None:
    _mcat = str(mbrow.get("Category", ""))
    with st.expander("Benchmark provision", expanded=False):
        st.markdown(str(mbrow.get("Provision", "")))

    _reg_opts    = _gap_reg_options(df_all, map_country, _mcat)
    _lbl_to_id   = {v: k for k, v in _reg_opts.items()}
    _default_lbls = [_reg_opts[e] for e in cur_eids if e in _reg_opts]

    _new_status = st.selectbox(
        "Compliance status", options=GAP_STATUS_OPTIONS,
        index=GAP_STATUS_OPTIONS.index(cur_status) if cur_status in GAP_STATUS_OPTIONS else 2,
    )
    if _reg_opts:
        _sel_lbls = st.multiselect(
            "Regulations that address this benchmark",
            options=list(_reg_opts.values()), default=_default_lbls,
        )
    else:
        st.info(
            f"No regulations found for {map_country}"
            + (f" under category '{_mcat}'" if _mcat else "")
            + ". Add regulations in the Editor tab first."
        )
        _sel_lbls = []

    _new_notes = st.text_area(
        "Gap assessment — what is still needed to fully meet this benchmark? (optional)",
        value=cur_notes if cur_notes not in ("", "nan") else "",
    )

    if st.button("Save mapping", type="primary", width="stretch"):
        _sel_eids = [_lbl_to_id[l] for l in _sel_lbls if l in _lbl_to_id]
        _now = pd.Timestamp.utcnow().isoformat()
        _new_mrow = {
            "Mapping_ID":     mapping_id,
            "Benchmark_ID":   bid,
            "Country":        map_country,
            "Status":         _new_status,
            "Entry_IDs":      ",".join(_sel_eids),
            "Gap_Assessment": _new_notes.strip(),
            "Updated_By":     auth_user,
            "Updated_At":     _now,
        }
        _before_m = (
            df_maps[(df_maps["Benchmark_ID"] == bid) & (df_maps["Country"] == map_country)]
            .iloc[0].to_dict() if not is_new_map else {}
        )
        if is_new_map:
            _upd_maps = pd.concat([df_maps, pd.DataFrame([_new_mrow])], ignore_index=True)
        else:
            _upd_maps = df_maps.copy()
            for _mi in _upd_maps.index[
                (_upd_maps["Benchmark_ID"] == bid) & (_upd_maps["Country"] == map_country)
            ].tolist():
                for _mk, _mv in _new_mrow.items():
                    _upd_maps.at[_mi, _mk] = _mv
        try:
            save_gap_mappings(_upd_maps, gap_dir)
            append_gap_audit(
                gap_dir, "add" if is_new_map else "edit", "mapping",
                f"{bid}:{map_country}", _before_m, _new_mrow, auth_user,
            )
            st.success("Mapping saved.")
            st.rerun()
        except Exception as _e:
            st.error(f"Save failed: {_e}")


# =========================
# Authentication — Keycloak integration point
# =========================
# Token flow:  ?token=<jwt>  (production: Keycloak-signed JWT)
# Legacy flow: ?country=X&user=Y  (fallback while migrating to token flow)
#
# Claims extracted from token:
#   preferred_username  — display name shown in UI / audit log
#   country             — ASEAN country name, or "NA" for admin
#   realm_roles         — list; "dashboard-admin" grants IS_ADMIN
#
# To plug in real Keycloak:
#   1. pip install python-jose[cryptography] requests
#   2. Set env vars KEYCLOAK_JWKS_URL and KEYCLOAK_CLIENT_ID
#   3. Replace the TODO block below with the jose.jwt.decode() call

_DUMMY_TOKENS: dict[str, dict] = {
    # Dummy tokens for development — remove or extend before production
    "dummy-admin": {
        "sub": "dev-admin-001",
        "preferred_username": "Admin",
        "country": "NA",
        "realm_roles": ["dashboard-admin", "dashboard-editor"],
    },
    "dummy-vietnam": {
        "sub": "dev-vn-001",
        "preferred_username": "VN_Editor",
        "country": "Viet Nam",
        "realm_roles": ["dashboard-editor"],
    },
    "dummy-singapore": {
        "sub": "dev-sg-001",
        "preferred_username": "SG_Editor",
        "country": "Singapore",
        "realm_roles": ["dashboard-editor"],
    },
}


def parse_auth_token(token: str) -> dict:
    """
    Parse and validate an auth token; return claims dict (empty = invalid/unauthenticated).

    Dummy tokens are accepted when KEYCLOAK_JWKS_URL is not configured.
    In production, replace the TODO block with real JWT validation.
    """
    import os
    token = token.strip()

    # ── Dummy tokens (development only) ─────────────────────────────────────
    if token in _DUMMY_TOKENS:
        return _DUMMY_TOKENS[token]

    # ── TODO: Real Keycloak JWT validation ───────────────────────────────────
    # jwks_url   = os.environ.get("KEYCLOAK_JWKS_URL", "")
    # client_id  = os.environ.get("KEYCLOAK_CLIENT_ID", "")
    # if jwks_url and client_id:
    #     try:
    #         from jose import jwt as _jwt, JWTError
    #         from jose.backends import RSAKey
    #         import requests, json as _json
    #         jwks = requests.get(jwks_url, timeout=5).json()
    #         claims = _jwt.decode(
    #             token, jwks, algorithms=["RS256"],
    #             audience=client_id, options={"verify_at_hash": False},
    #         )
    #         return claims
    #     except Exception:
    #         return {}
    # ─────────────────────────────────────────────────────────────────────────

    return {}  # Unknown token → public view


def _sanitize(s: str, pattern: str, maxlen: int = 50) -> Optional[str]:
    return re.sub(pattern, "", s)[:maxlen] or None


# =========================
# Auth resolution (must happen before tabs are created)
# =========================
_qp = get_query_params()

# Token-based auth (Keycloak-ready)
_token  = _qp.get("token", "").strip()
_claims = parse_auth_token(_token) if _token else {}

if _claims:
    # Token path — derive everything from claims
    _auth_country = _sanitize(_claims.get("country") or "", r"[^A-Za-z0-9 \-_]")
    _auth_user    = _sanitize(_claims.get("preferred_username") or "", r"[^A-Za-z0-9 \-_@.]")
    _auth_roles: list[str] = _claims.get("realm_roles", [])
    IS_ADMIN         = "dashboard-admin" in _auth_roles
    IS_AUTHENTICATED = True
else:
    # Legacy fallback: ?country=X&user=Y (used until token flow is live)
    _auth_country = _sanitize(_qp.get("country") or "", r"[^A-Za-z0-9 \-_]")
    _auth_user    = _sanitize(_qp.get("user") or "", r"[^A-Za-z0-9 \-_@.]")
    _auth_roles   = []
    IS_ADMIN         = (_auth_country == "NA" and _auth_user == "Admin")
    IS_AUTHENTICATED = bool(_auth_country and _auth_user)

# =========================
# Navigation
# =========================
_NAV_ICONS = {"Map": "🗺️", "Table": "📋", "Gap Analysis": "📊", "Editor": "✏️"}
_nav_pages = ["Map", "Table", "Gap Analysis", "Editor"] if IS_AUTHENTICATED else ["Map", "Table"]

if "active_page" not in st.session_state:
    _qp_page = _qp.get("page", "Map")
    st.session_state["active_page"] = _qp_page if _qp_page in _nav_pages else "Map"
_active_page = st.session_state.get("active_page", "Map")
if _active_page not in _nav_pages:
    _active_page = "Map"
    st.session_state["active_page"] = "Map"


def _nav_to(page: str) -> None:
    st.session_state["active_page"] = page
    try:
        st.query_params.update({**dict(st.query_params), "page": page})
    except Exception:
        pass
    st.rerun()


# ── Horizontal top nav bar ────────────────────────────────────────────────
_NAV_HELP: dict[str, list[tuple[str, str]]] = {
    "Map": [
        ("Reading the map",
         "**Highlighted countries** have at least one regulation in the current filter.\n\n"
         "**Hover** over a country to see a tooltip with regulation count and recent titles.\n\n"
         "Use the **country selector** below the map to open a full detail panel with regulator(s), "
         "full regulation list, and key provisions by category."),
        ("Exports",
         "- **All categories, no country** → Excel, one sheet per category\n"
         "- **Single category, no country** → CSV\n"
         "- **Country selected** → Excel + Print/PDF inside the country panel; top-right export hidden"),
    ],
    "Table": [
        ("Reading the table",
         "**Category = All:** Matrix — one row per country, ✓ where a regulation exists.\n\n"
         "**Specific category:** Full worksheet rows, grouped by country.\n\n"
         "Click any row to open a country detail panel. Select two rows for a side-by-side comparison."),
        ("Exports",
         "The **Export** button exports the currently displayed table to CSV, "
         "respecting active category and country filters."),
    ],
    "Gap Analysis": [
        ("Overview",
         "Shows all benchmarks grouped by Standard. Click a row to open a benchmark detail popup "
         "with per-country status and linked regulations. Use **Export** for a flat CSV."),
        ("Benchmarks",
         "Admins can add, edit, and delete benchmark definitions. "
         "Each benchmark has a Standard, Category, Topic, and Provision text. Use **Export** for CSV."),
        ("Mappings",
         "Select a country to see benchmarks grouped by **Category → Standard**. "
         "Click a row to view status, linked regulations, and gap assessment. "
         "Use the **search bar** to filter. **Export** downloads the full gap analysis as Excel."),
    ],
    "Editor": [
        ("Adding a regulation",
         "Click **+ Add regulation**, fill in Category, Regulator, Year, Title, Source URL, then save."),
        ("Editing / Archiving",
         "Expand a category, click a row — an edit dialog opens. "
         "Save changes or archive (permanently removes from live data, preserved in audit log)."),
        ("Audit log",
         "**Change history** button at the bottom shows the 50 most recent changes with before/after snapshots."),
    ],
}

_nc      = len(_nav_pages)
_nav_col_widths = [0.9] + [1] * _nc + [0.18]
_nav_cols       = st.columns(_nav_col_widths, vertical_alignment="center")

# Title + user badge stacked in col 0
if IS_AUTHENTICATED:
    _badge_label = f"{_auth_user} · admin access" if IS_ADMIN else f"{_auth_user} · {_auth_country}"
    _badge_html  = (
        f"<br><span style='font-size:0.62rem;color:#64748b;white-space:nowrap'>"
        f"👤 {_badge_label}</span>"
        f"<br><span style='font-size:0.58rem;color:#334155;white-space:nowrap'>"
        f"AMS-exclusive</span>"
    )
else:
    _badge_html = ""
_nav_cols[0].markdown(
    f"<span style='font-size:0.88rem;font-weight:700;color:#e2e8f0;white-space:nowrap'>"
    f"<abbr title='ASEAN Regulatory Information System' style='text-decoration:none;cursor:default'>ARIS</abbr></span>"
    f"<span style='font-size:0.58rem;color:#475569;margin-left:5px'>v{_APP_VERSION}</span>"
    f"{_badge_html}",
    unsafe_allow_html=True,
)

# Nav buttons
for _ni, _page in enumerate(_nav_pages):
    if _nav_cols[_ni + 1].button(
        f"{_NAV_ICONS.get(_page, '')} {_page}",
        key=f"nav_{_page}",
        type="primary" if _active_page == _page else "secondary",
        width="stretch",
    ):
        _nav_to(_page)

# Help popover — last column
with _nav_cols[-1].popover("❓", help="Page help"):
    _help_items = _NAV_HELP.get(_active_page, [])
    if _help_items:
        for _htitle, _hbody in _help_items:
            with st.expander(f"**{_htitle}**", expanded=False):
                st.markdown(_hbody)
    else:
        st.caption("No help available for this page.")

def _clear_editor_table_keys() -> None:
    for _k in list(st.session_state.keys()):
        if _k.startswith("editor_table_"):
            del st.session_state[_k]


@st.dialog("Change history", width="large")
def _dlg_change_history(df: pd.DataFrame, col_cfg: dict) -> None:
    if df.empty:
        st.info("No changes recorded yet.")
    else:
        st.dataframe(df, width="stretch", hide_index=True, height=520, column_config=col_cfg)


# =========================
# MAP PAGE
# =========================
if _active_page == "Map":
    # ── Header row: Category | Country | [space] | Export ──
    _cur_cat     = st.session_state.get(_CAT_KEY_MAP) or "All"
    _cur_cat_idx = categories.index(_cur_cat) if _cur_cat in categories else 0
    _mc1, _mc2, _, _mc_dl = st.columns([1, 1, 2, 1], vertical_alignment="bottom")

    sel_category = _mc1.selectbox(
        "Category", options=categories, index=_cur_cat_idx,
        key=_CAT_KEY_MAP, on_change=_on_map_category_change,
    )
    df_f = df_all[df_all["Category"] == sel_category].copy() if sel_category != "All" else df_all.copy()

    # Build country list for the selectbox
    by_country = (
        df_f.groupby("Country_std", dropna=True)
        .size()
        .reset_index(name="Regulation_Count")
        .rename(columns={"Country_std": "Country"})
    )
    hover_texts = []
    for c in by_country["Country"].tolist():
        latest10 = latest_regs_by_country(df_f, c, n=10)
        hover_texts.append(build_hover_list(latest10))
    by_country["Latest_10"] = hover_texts
    by_country["ISO3"] = by_country["Country"].map(ASEAN_ISO3)

    _map_country_opts = sorted(
        c for c in by_country["Country"].tolist() if pd.notna(c) and str(c) != ""
    )
    map_country = _mc2.selectbox(
        "Country", options=_map_country_opts, index=None,
        placeholder="Select a country…", key="map_country_sel",
    )
    _country_selected = map_country is not None

    if not _country_selected:
        if sel_category == "All":
            _mc_dl.download_button(
                "⬇ Export (.xlsx)", _build_regulations_excel(df_f, by_category=True),
                "asean_regulations_all.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                width="stretch", key="dl_export_map",
            )
        else:
            _mc_dl.download_button(
                "⬇ Export (.csv)",
                df_f[[c for c in df_f.columns if c not in _REG_EXPORT_EXCLUDE]]
                .rename(columns=_REG_EXPORT_RENAME)
                .to_csv(index=False).encode("utf-8"),
                f"asean_{sel_category.lower().replace(' ', '_')}.csv", "text/csv",
                width="stretch", key="dl_export_map",
            )

    # ── Single ASEAN map — selected country highlighted gold, others blue ──
    by_country["Highlight"] = by_country["Country"].apply(
        lambda c: 1 if (_country_selected and c == map_country) else 0
    )

    _map_cfg = {"displayModeBar": False, "scrollZoom": False, "doubleClick": False, "responsive": True}

    _asean_fig = px.choropleth(
        by_country,
        locations="ISO3",
        locationmode="ISO-3",
        color="Highlight",
        color_continuous_scale=["#2563eb", "#facc15"],
        range_color=(0, 1),
        custom_data=["Country", "Regulation_Count", "Latest_10"],
    )
    _asean_fig.update_traces(
        hovertemplate="<b>%{customdata[0]}</b><br><br>%{customdata[2]}<extra></extra>",
        showscale=False,
    )
    _asean_fig.update_layout(
        title_text="", hoverlabel=dict(align="left"),
        dragmode=False, hovermode="closest", coloraxis_showscale=False,
        autosize=False, width=None, height=520,
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=10, b=0),
    )
    _asean_fig.update_geos(
        scope="asia", projection_type="mercator",
        lonaxis=dict(range=[92, 141]), lataxis=dict(range=[-11, 24]),
        visible=True,
        showcoastlines=True, coastlinecolor="rgba(255,255,255,0.35)",
        showcountries=True, countrycolor="rgba(255,255,255,0.85)",
        showland=True, landcolor="rgba(20, 30, 45, 1)",
        showocean=True, oceancolor="rgba(10, 16, 26, 1)",
        showlakes=True, lakecolor="rgba(10, 16, 26, 1)",
        bgcolor="rgba(0,0,0,0)",
    )
    _asean_fig.layout.width = None

    st.plotly_chart(_asean_fig, width="stretch", key="asean_main_map", config=_map_cfg)

    if by_country.empty:
        st.info("No regulations found for the current category filter.")
    elif not _country_selected:
        st.caption("Hover a country to preview its 10 most recent regulations (based on the current filters).")
    else:
        st.caption(f"Showing details for {map_country} — scroll down for more.")
        show_country_modal(map_country, df_f, key_suffix="map", show_minimap=False, show_pdf_btn=True)

# =========================
# TABLE PAGE
# =========================
elif _active_page == "Table":
    sel_category = _render_category_filter("table")
    df_f = df_all[df_all["Category"] == sel_category].copy() if sel_category != "All" else df_all.copy()

    all_sheet_names = sorted(df_all["Category"].dropna().unique().tolist())

    # =========================================================
    # MODE A: Category = All -> summary matrix
    # =========================================================
    if sel_category == "All":
        regs_by_country = (
            df_f.groupby("Country_std")["Regulator_std"]
            .apply(
                lambda x: "\n".join(sorted({str(v) for v in x if pd.notna(v)}))
            )
            .reset_index()
            .rename(columns={"Country_std": "Country", "Regulator_std": "Regulator"})
        )

        counts = (
            df_f.groupby(["Country_std", "Category"])
            .size()
            .reset_index(name="Count")
            .assign(HasReg=True)
            .pivot(index="Country_std", columns="Category", values="HasReg")
            .fillna(False)
            .astype(bool)
            .reset_index()
            .rename(columns={"Country_std": "Country"})
        )

        t = regs_by_country.merge(counts, on="Country", how="outer").fillna({"Regulator": ""})
        for s in all_sheet_names:
            if s not in t.columns:
                t[s] = False

        t["Country"] = t["Country"].fillna("").astype(str)
        t.insert(0, "Flag", t["Country"].map(country_flag))
        t = t[["Flag", "Country"] + all_sheet_names].sort_values("Country")

        CHECK, BLANK = "✓", ""
        for s in all_sheet_names:
            t[s] = t[s].map(lambda x: CHECK if x else BLANK)

        country_options = sorted(c for c in t["Country"].tolist() if c)
        a_sel = st.multiselect(
            "Select countries to compare",
            options=country_options,
            placeholder="Choose one or more countries…",
            key="table_a_country_sel",
        )
        _render_data_table(t, highlighted=set(a_sel), row_height=36)
        if a_sel:
            if len(a_sel) == 1:
                show_country_modal(a_sel[0], df_f, key_suffix="table_0")
            else:
                show_country_comparison(a_sel, df_f, key_suffix="table_cmp")

    # =========================================================
    # MODE B: Category selected -> country picker then provisions
    # =========================================================
    else:
        d = df_f.copy()

        known_meta_b = _KNOWN_META
        detail_cols_b = [
            c for c in d.columns
            if c not in known_meta_b
            and d[c].dropna().astype(str).str.strip()
               .pipe(lambda s: s[~s.str.lower().isin({"nan", "none", ""})]).any()
        ]

        country_options = sorted(
            c for c in d["Country_std"].dropna().unique() if str(c).strip()
        )
        b_sel = st.multiselect(
            "Select a country to view",
            options=country_options,
            placeholder="Choose one or more countries…",
            key="table_b_country_sel",
        )

        if not b_sel:
            # Overview matrix: rows = countries, cols = key provision fields
            if detail_cols_b:
                CHECK, BLANK = "✓", ""
                overview_rows = []
                for country in country_options:
                    dc = d[d["Country_std"] == country]
                    row: dict = {"Flag": country_flag(country), "Country": country}
                    for col in detail_cols_b:
                        if col not in dc.columns:
                            row[col] = BLANK
                            continue
                        has_val = (
                            dc[col].dropna().astype(str).str.strip()
                            .pipe(lambda s: s[~s.str.lower().isin({"nan", "none", ""})])
                            .any()
                        )
                        row[col] = CHECK if has_val else BLANK
                    overview_rows.append(row)
                if overview_rows:
                    overview_df = pd.DataFrame(overview_rows)
                    _render_data_table(overview_df, row_height=36)
                else:
                    st.caption("No data available for this category.")
            else:
                st.caption("No key provision fields found for this category.")
        else:
            # Header line — flags + names
            header = "  ·  ".join(f"{country_flag(c)} {c}" for c in b_sel)
            st.markdown(f"## {header}")

            # Build one row per provision field, one column per country
            col_headers = ["Field"] + [f"{country_flag(c)} {c}" for c in b_sel]
            rows = []
            for col in detail_cols_b:
                if col not in d.columns:
                    continue
                row: dict = {"Field": col}
                has_any = False
                for country in b_sel:
                    dc = d[d["Country_std"] == country]
                    vals = (
                        dc[col].dropna().astype(str).str.strip()
                        .pipe(lambda s: s[~s.str.lower().isin({"nan", "none", ""})])
                        .unique().tolist()
                    )
                    cell = "\n".join(sorted(vals))
                    row[f"{country_flag(country)} {country}"] = cell
                    if cell:
                        has_any = True
                if has_any:
                    rows.append(row)

            if rows:
                _render_provision_table(rows, col_headers)
            else:
                st.caption("No provisions data available for the selected countries and category.")


# =========================
# GAP ANALYSIS PAGE  (IS_AUTHENTICATED only)
# =========================
elif _active_page == "Gap Analysis" and IS_AUTHENTICATED:
    if True:
        _gap_dir     = resolved_data_dir.parent          # src/ not src/categories/
        _gap_cache   = get_gap_cache_key(_gap_dir)
        _df_bench    = load_gap_benchmarks(str(_gap_dir), _gap_cache)
        _df_maps     = load_gap_mappings(str(_gap_dir), _gap_cache)
        _gap_cats    = sorted(df_all["Category"].dropna().unique().tolist())

        # Gap country: admin can select any; country user is locked to their URL param
        _gap_country: Optional[str] = None
        if IS_ADMIN:
            _gap_country = "Admin"  # Admin sees all; per-country selection handled inside Mappings tab
        else:
            _gap_country = _auth_country

        # Session-state-backed sub-tabs so rerun doesn't reset to Overview
        if "gap_subtab" not in st.session_state:
            st.session_state["gap_subtab"] = "Overview"
        _gsub_names = ["Overview", "Benchmarks", "Mappings"]
        st.markdown("<div style='border-bottom:1px solid #1e293b;margin-bottom:8px'></div>", unsafe_allow_html=True)
        _gsub_cols  = st.columns(len(_gsub_names))
        for _gi, _gn in enumerate(_gsub_names):
            if _gsub_cols[_gi].button(
                _gn, key=f"gap_stab_{_gn}", width="stretch",
                type="primary" if st.session_state["gap_subtab"] == _gn else "secondary",
            ):
                st.session_state["gap_subtab"] = _gn
                st.rerun()
        _active_gsub = st.session_state["gap_subtab"]
        st.divider()

        # ── Overview ──────────────────────────────────────────────
        if _active_gsub == "Overview":
            # Legend
            st.markdown(
                "<div style='display:flex;gap:8px;flex-wrap:wrap;margin-bottom:14px'>"
                "<span style='background:#22c55e22;border:1px solid #22c55e55;border-radius:12px;"
                "padding:3px 11px;font-size:0.8rem;color:#22c55e'>✅ Meets</span>"
                "<span style='background:#f59e0b22;border:1px solid #f59e0b55;border-radius:12px;"
                "padding:3px 11px;font-size:0.8rem;color:#f59e0b'>⚠️ Partially meets</span>"
                "<span style='background:#ef444422;border:1px solid #ef444455;border-radius:12px;"
                "padding:3px 11px;font-size:0.8rem;color:#ef4444'>✗ Does not meet</span>"
                "<span style='background:#64748b22;border:1px solid #64748b55;border-radius:12px;"
                "padding:3px 11px;font-size:0.8rem;color:#64748b'>— Not assessed</span>"
                "<span style='margin-left:6px;font-size:0.78rem;color:#475569;"
                "align-self:center'>Click a row to view details</span>"
                "</div>",
                unsafe_allow_html=True,
            )

            if _df_bench.empty:
                st.info("No benchmarks have been defined yet. Admin can add them in the Benchmarks tab.")
            else:
                # Build flat CSV: one row per benchmark, Standard as column, worded status
                _ov_csv_rows = []
                for _, _br in _df_bench.iterrows():
                    _bid = str(_br["Benchmark_ID"])
                    _crow: dict = {
                        "Standard":     str(_br.get("Standard", "")),
                        "Benchmark_ID": _bid,
                        "Topic":        str(_br.get("Topic", "")),
                        "Category":     str(_br.get("Category", "")),
                    }
                    for _c in ASEAN_COUNTRIES_ORDERED:
                        _cm = _df_maps[
                            (_df_maps["Benchmark_ID"] == _bid) &
                            (_df_maps["Country"] == _c)
                        ]
                        _crow[_c] = (
                            "Not assessed" if _cm.empty
                            else str(_cm.iloc[0].get("Status", "Not assessed"))
                        )
                    _ov_csv_rows.append(_crow)
                _ov_export_csv = pd.DataFrame(_ov_csv_rows).to_csv(index=False).encode("utf-8")
                _, _ov_dl_col = st.columns([5, 1])
                _ov_dl_col.download_button(
                    "⬇ Export", _ov_export_csv,
                    "gap_analysis_overview.csv", "text/csv",
                    width="stretch", key="gap_ov_export",
                )

                _ov_standards = (
                    _df_bench["Standard"].dropna().unique().tolist()
                    if "Standard" in _df_bench.columns else [""]
                )
                for _ov_si, _ov_std in enumerate(_ov_standards):
                    _ov_bench = _df_bench[_df_bench["Standard"] == _ov_std].copy()
                    _ov_rows  = []
                    for _, _br in _ov_bench.iterrows():
                        _bid = str(_br["Benchmark_ID"])
                        _row: dict = {
                            "Topic":    str(_br.get("Topic", "")),
                            "Category": str(_br.get("Category", "")),
                        }
                        for _c in ASEAN_COUNTRIES_ORDERED:
                            _cm = _df_maps[
                                (_df_maps["Benchmark_ID"] == _bid) &
                                (_df_maps["Country"] == _c)
                            ]
                            _st = (
                                "Not assessed" if _cm.empty
                                else str(_cm.iloc[0].get("Status", "Not assessed"))
                            )
                            _row[country_flag(_c)] = GAP_STATUS_META.get(_st, GAP_STATUS_META["Not assessed"])["emoji"]
                        _ov_rows.append(_row)

                    _ov_df = pd.DataFrame(_ov_rows)
                    _flag_cols = [country_flag(_c) for _c in ASEAN_COUNTRIES_ORDERED]
                    _ov_col_cfg: dict = {
                        "Topic":    st.column_config.TextColumn("Topic / Area", width="large"),
                        "Category": st.column_config.TextColumn("Category",     width="small"),
                    }
                    for _fc in _flag_cols:
                        _ov_col_cfg[_fc] = st.column_config.TextColumn(_fc, width="small")

                    with st.expander(f"**{_ov_std}**", expanded=True):
                        _ov_event = st.dataframe(
                            _ov_df, width="stretch", hide_index=True,
                            height=min(400, max(80, 40 + len(_ov_rows) * 36)),
                            key=f"gap_ov_tbl_{_ov_si}",
                            on_select="rerun",
                            selection_mode="single-row",
                            column_config=_ov_col_cfg,
                        )
                        _ov_sel = (
                            _ov_event.selection.rows
                            if hasattr(_ov_event, "selection") and _ov_event.selection
                            else []
                        )
                        if _ov_sel:
                            _dlg_benchmark_detail(
                                _ov_bench.iloc[_ov_sel[0]],
                                _df_maps, df_all,
                            )

        # ── Benchmarks ────────────────────────────────────────────
        elif _active_gsub == "Benchmarks":
            if IS_ADMIN:
                _bcol1, _bcol2, _bcol3 = st.columns([3, 1, 1])
                _bcol1.subheader("Benchmark definitions")
                if _bcol2.button("＋ Add benchmark", width="stretch", type="primary"):
                    _dlg_add_benchmark(_df_bench, _gap_cats, _gap_dir, _auth_user or "Admin")
            else:
                _, _bcol2, _bcol3 = st.columns([3, 1, 1])

            if _df_bench.empty:
                st.info("No benchmarks defined yet.")
            else:
                _bench_export_df = _df_bench[
                    ["Benchmark_ID", "Standard", "Category", "Topic", "Provision"]
                ].rename(columns={"Benchmark_ID": "Benchmark ID"})
                _bcol3.download_button(
                    "⬇ Export", _bench_export_df.to_csv(index=False).encode("utf-8"),
                    "benchmark_definitions.csv", "text/csv",
                    width="stretch", key="gap_bench_export",
                )

                _display_bench = _df_bench[
                    ["Benchmark_ID", "Standard", "Category", "Topic", "Provision",
                     "Updated_By", "Updated_At"]
                ].copy()
                _bench_event = st.dataframe(
                    _display_bench,
                    width="stretch",
                    hide_index=True,
                    height=min(500, max(120, 40 + len(_df_bench) * 36)),
                    key="gap_bench_table",
                    on_select="rerun" if IS_ADMIN else "ignore",
                    selection_mode="single-row",
                    column_config={
                        "Benchmark_ID": None,
                        "Standard":     st.column_config.TextColumn("Standard",    width="medium"),
                        "Category":     st.column_config.TextColumn("Category",    width="small"),
                        "Topic":        st.column_config.TextColumn("Topic / Area", width="large"),
                        "Provision":    st.column_config.TextColumn("Provision",   width="large"),
                        "Updated_By":   None,
                        "Updated_At":   None,
                    },
                )
                if IS_ADMIN:
                    _bsel_idx = (
                        _bench_event.selection.rows
                        if hasattr(_bench_event, "selection") and _bench_event.selection
                        else []
                    )
                    if _bsel_idx:
                        _bsel_row = _df_bench.iloc[_bsel_idx[0]]
                        _bid_sel  = str(_bsel_row["Benchmark_ID"])
                        _bet_col1, _bet_col2 = st.columns(2)
                        if _bet_col1.button("✏️ Edit", key="gap_bench_edit_btn", width="stretch"):
                            _dlg_edit_benchmark(
                                _bsel_row, _df_bench, _gap_cats, _gap_dir, _auth_user or "Admin"
                            )
                        # Delete — only if no mappings reference it
                        _ref_count = len(_df_maps[_df_maps["Benchmark_ID"] == _bid_sel])
                        if _ref_count > 0:
                            _bet_col2.button(
                                f"🗑 Delete (blocked)",
                                disabled=True,
                                width="stretch",
                                help=f"This benchmark has {_ref_count} country mapping(s). Remove all mappings for '{_bid_sel}' in the Mappings tab before deleting.",
                            )
                        else:
                            _del_key = f"gap_bench_del_confirm_{_bid_sel}"
                            if not st.session_state.get(_del_key):
                                if _bet_col2.button("🗑 Delete", width="stretch", type="secondary",
                                                    key="gap_bench_del_btn"):
                                    st.session_state[_del_key] = True
                                    st.rerun()
                            else:
                                st.warning(f"Delete **{_bid_sel}**? This cannot be undone.")
                                _dc1, _dc2 = st.columns(2)
                                if _dc1.button("Yes, delete", type="secondary", width="stretch", key="gap_bench_del_yes"):
                                    _before_del = _bsel_row.to_dict()
                                    _updated_bench = _df_bench[_df_bench["Benchmark_ID"] != _bid_sel].reset_index(drop=True)
                                    try:
                                        save_gap_benchmarks(_updated_bench, _gap_dir)
                                        append_gap_audit(
                                            _gap_dir, "delete", "benchmark",
                                            _bid_sel, _before_del, {}, _auth_user or "Admin"
                                        )
                                    except Exception as _e:
                                        st.error(f"Delete failed: {_e}")
                                    st.session_state.pop(_del_key, None)
                                    st.rerun()
                                if _dc2.button("Cancel", type="primary", width="stretch", key="gap_bench_del_no"):
                                    st.session_state.pop(_del_key, None)
                                    st.rerun()

        # ── Mappings ──────────────────────────────────────────────
        elif _active_gsub == "Mappings":
            if IS_ADMIN:
                _map_country = st.selectbox(
                    "Country", options=ASEAN_COUNTRIES_ORDERED, index=None,
                    placeholder="Select a country…", key="gap_map_country_sel"
                )
            else:
                _map_country = _auth_country or ""
                st.caption(f"Mapping regulations for: {country_flag(_map_country)} **{_map_country}**")

            if _df_bench.empty:
                st.info("No benchmarks available. Ask an admin to add benchmarks first.")
            elif _map_country:
                # ── Header: search + export ────────────────────────
                _mh1, _mh2 = st.columns([4, 1])
                _msearch = _mh1.text_input(
                    "Search", placeholder="Search standard, benchmark or topic…",
                    label_visibility="collapsed", key="gap_map_search",
                ).strip().lower()

                # Build export — all benchmarks × country (flat Excel sheet)
                _map_exp_rows: list[dict] = []
                with st.spinner("Building export…"):
                    for _, _ebr in _df_bench.iterrows():
                        _ebid = str(_ebr["Benchmark_ID"])
                        _ecm  = _df_maps[
                            (_df_maps["Benchmark_ID"] == _ebid) &
                            (_df_maps["Country"] == _map_country)
                        ]
                        _est  = "Not assessed" if _ecm.empty else str(_ecm.iloc[0].get("Status", "Not assessed"))
                        _eeids = (
                            [e.strip() for e in str(_ecm.iloc[0].get("Entry_IDs", "")).split(",") if e.strip()]
                            if not _ecm.empty else []
                        )
                        _elinked = df_all[df_all["Entry_ID"].isin(_eeids)] if _eeids else pd.DataFrame()
                        _eregs = "; ".join(
                            str(_rr.get("Regulation_Title", ""))
                            + (f" ({format_year(_rr.get('Year', ''))})"
                               if str(_rr.get("Year","")) not in ("","nan") else "")
                            for _, _rr in _elinked.iterrows()
                        )
                        _enotes = str(_ecm.iloc[0].get("Gap_Assessment", "") if not _ecm.empty else "").strip()
                        if _enotes in ("nan",): _enotes = ""
                        _map_exp_rows.append({
                            "Category":           str(_ebr.get("Category", "")),
                            "Standard":           str(_ebr.get("Standard", "")),
                            "Benchmark ID":       _ebid,
                            "Topic":              str(_ebr.get("Topic", "")),
                            "Provision":          str(_ebr.get("Provision", "")),
                            "Status":             _est,
                            "Linked Regulations": _eregs,
                            "Gap Assessment":     _enotes,
                        })
                    _map_exp_buf = io.BytesIO()
                    with pd.ExcelWriter(_map_exp_buf, engine="openpyxl") as _ew:
                        pd.DataFrame(_map_exp_rows).to_excel(
                            _ew, sheet_name=f"{_map_country[:31]}", index=False
                        )
                _mh2.download_button(
                    "⬇ Export", _map_exp_buf.getvalue(),
                    f"gap_{_map_country.lower().replace(' ','_')}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    width="stretch", key="gap_map_export",
                )

                # ── Filter by search ───────────────────────────────
                _map_bench_f = _df_bench.copy()
                if _msearch:
                    _map_bench_f = _map_bench_f[
                        _map_bench_f["Standard"].fillna("").str.lower().str.contains(_msearch, regex=False)
                        | _map_bench_f["Benchmark_ID"].fillna("").str.lower().str.contains(_msearch, regex=False)
                        | _map_bench_f["Topic"].fillna("").str.lower().str.contains(_msearch, regex=False)
                    ]

                if _map_bench_f.empty:
                    st.info("No benchmarks match your search.")
                else:
                    _can_edit = IS_ADMIN or (_auth_country == _map_country)
                    _tbl_idx  = 0

                    for _mcat in sorted(_map_bench_f["Category"].dropna().unique()):
                        _mcat_bench = _map_bench_f[_map_bench_f["Category"] == _mcat]
                        with st.expander(f"**{_mcat}**", expanded=True):
                            for _mstd in sorted(_mcat_bench["Standard"].dropna().unique()):
                                _mstd_bench = _mcat_bench[_mcat_bench["Standard"] == _mstd].copy()
                                st.markdown(
                                    f"<div style='font-size:0.78rem;font-weight:600;"
                                    f"color:#94a3b8;text-transform:uppercase;"
                                    f"letter-spacing:0.06em;margin:10px 0 4px'>{_mstd}</div>",
                                    unsafe_allow_html=True,
                                )

                                # Build table rows
                                _mtbl_rows: list[dict] = []
                                for _, _mbr in _mstd_bench.iterrows():
                                    _mbid = str(_mbr["Benchmark_ID"])
                                    _mcm  = _df_maps[
                                        (_df_maps["Benchmark_ID"] == _mbid) &
                                        (_df_maps["Country"] == _map_country)
                                    ]
                                    _mst = "Not assessed" if _mcm.empty else str(_mcm.iloc[0].get("Status", "Not assessed"))
                                    _mtbl_rows.append({
                                        "_bid":   _mbid,
                                        "Topic":  str(_mbr.get("Topic", "")),
                                        "Status": GAP_STATUS_META.get(_mst, GAP_STATUS_META["Not assessed"])["emoji"] + " " + _mst,
                                    })

                                _mtbl_df    = pd.DataFrame(_mtbl_rows)
                                _mtbl_event = st.dataframe(
                                    _mtbl_df.drop(columns=["_bid"]),
                                    width="stretch", hide_index=True,
                                    height=min(300, max(56, 40 + len(_mtbl_rows) * 35)),
                                    key=f"gap_map_t_{_tbl_idx}",
                                    on_select="rerun",
                                    selection_mode="single-row",
                                    column_config={
                                        "Topic":  st.column_config.TextColumn("Topic / Area", width="large"),
                                        "Status": st.column_config.TextColumn("Status",       width="medium"),
                                    },
                                )
                                _mtbl_sel = (
                                    _mtbl_event.selection.rows
                                    if hasattr(_mtbl_event, "selection") and _mtbl_event.selection
                                    else []
                                )

                                if _mtbl_sel:
                                    _sel_idx  = _mtbl_sel[0]
                                    _sel_mbid = _mtbl_rows[_sel_idx]["_bid"]
                                    _mbrow    = _mstd_bench[_mstd_bench["Benchmark_ID"] == _sel_mbid].iloc[0]

                                    with st.expander("Benchmark provision", expanded=False):
                                        st.markdown(str(_mbrow.get("Provision", "")))

                                    _exist_map = _df_maps[
                                        (_df_maps["Benchmark_ID"] == _sel_mbid) &
                                        (_df_maps["Country"] == _map_country)
                                    ]
                                    if _exist_map.empty:
                                        _cur_status = "Not assessed"
                                        _cur_eids: list[str] = []
                                        _cur_notes  = ""
                                        _mapping_id = str(uuid.uuid4())
                                        _is_new_map = True
                                        _em         = None
                                    else:
                                        _em         = _exist_map.iloc[0]
                                        _cur_status = str(_em.get("Status", "Not assessed"))
                                        _cur_eids   = [e.strip() for e in str(_em.get("Entry_IDs", "")).split(",") if e.strip()]
                                        _cur_notes  = str(_em.get("Gap_Assessment", _em.get("Notes", "")))
                                        if _cur_notes in ("nan",): _cur_notes = ""
                                        _mapping_id = str(_em.get("Mapping_ID", uuid.uuid4()))
                                        _is_new_map = False

                                    _render_mapping_card(
                                        _cur_status, _cur_eids, _cur_notes,
                                        _is_new_map, _em, df_all,
                                    )

                                    if _can_edit:
                                        if st.button(
                                            "✏️ Edit mapping", type="primary",
                                            key=f"gap_map_edit_{_tbl_idx}",
                                        ):
                                            _dlg_edit_mapping(
                                                bid=_sel_mbid, map_country=_map_country,
                                                mbrow=_mbrow, cur_status=_cur_status,
                                                cur_eids=_cur_eids, cur_notes=_cur_notes,
                                                mapping_id=_mapping_id, is_new_map=_is_new_map,
                                                df_all=df_all, df_maps=_df_maps,
                                                gap_dir=_gap_dir, auth_user=_auth_user or "unknown",
                                            )

                                _tbl_idx += 1

        # ── Audit ledger ───────────────────────────────────────────
        st.divider()
        _gap_audit_path = _gap_dir / _GAP_AUDIT_FILE
        _gap_hist_df = pd.DataFrame()
        if _gap_audit_path.exists():
            _gap_hist_df = pd.read_csv(_gap_audit_path, dtype=str).fillna("")
            _gap_hist_df["Timestamp"] = pd.to_datetime(
                _gap_hist_df["Timestamp"], errors="coerce"
            )
            if not IS_ADMIN:
                def _audit_country(row: pd.Series) -> str:
                    if str(row.get("Record_Type", "")) == "mapping":
                        rid = str(row.get("Record_ID", ""))
                        return rid.split(":")[-1] if ":" in rid else ""
                    return ""
                _gap_hist_df = _gap_hist_df[
                    _gap_hist_df.apply(_audit_country, axis=1) == (_auth_country or "")
                ]
            _gap_hist_df = (
                _gap_hist_df.sort_values("Timestamp", ascending=False)
                .head(50)[["Timestamp", "User", "Action", "Record_Type", "Record_ID"]]
            )
        _gap_hist_col_cfg = {
            "Timestamp":   st.column_config.DatetimeColumn("When", format="YYYY-MM-DD HH:mm"),
            "User":        st.column_config.TextColumn("By",     width="small"),
            "Action":      st.column_config.TextColumn("Action", width="small"),
            "Record_Type": st.column_config.TextColumn("Type",   width="small"),
            "Record_ID":   st.column_config.TextColumn("Record", width="medium"),
        }
        if _active_gsub != "Overview":
            if st.button("📋 Change history", type="secondary", key="gap_hist_btn"):
                _dlg_change_history(_gap_hist_df, _gap_hist_col_cfg)


# =========================
# EDITOR PAGE  (only rendered when IS_AUTHENTICATED)
# =========================
elif _active_page == "Editor" and IS_AUTHENTICATED:
    if True:
        st.subheader("Country editor & audit log")

        editor_countries = sorted(df_all["Country_std"].dropna().unique().tolist())

        # ── Determine editor_country and display session banner ──────────
        if IS_ADMIN:
            # Admin can switch between any country via a dropdown
            st.success(f"Signed in as **{_auth_user}** · admin access to all countries")
            editor_country: Optional[str] = st.selectbox(
                "Select country to edit",
                options=["(Choose a country)"] + editor_countries,
                index=0,
                key="admin_country_select",
            )
            if editor_country == "(Choose a country)":
                editor_country = None
        else:
            # Country user — locked to their own country
            if _auth_country not in editor_countries:
                st.error(
                    f"Country **{_auth_country}** was not found in the dataset. "
                    "Please contact the administrator."
                )
                st.stop()
            editor_country = _auth_country
            st.success(
                f"Signed in as **{_auth_user}** · editing "
                f"{country_flag(editor_country or '')} **{editor_country}**"
            )

        if not editor_country:
            st.info("Choose a country above to view, edit, or add regulations.")
        else:
            country_rows = df_all[df_all["Country_std"] == editor_country].copy()
            country_rows = country_rows.sort_values(
                ["Year", "Regulation_Title"], ascending=[False, True]
            )

            category_choices = sorted(df_all["Category"].dropna().unique().tolist())

            regulator_name = (
                country_rows["Regulator_std"].dropna().astype(str).str.strip()
                .loc[lambda s: s != ""].iloc[0]
                if not country_rows["Regulator_std"].dropna().empty
                else "current records"
            )
            st.markdown(
                f"#### {country_flag(editor_country or '')} **{editor_country}** — {regulator_name}"
            )

            preview_cols = [
                "Year", "Regulation_Title", "Source_URL",
            ]
            preview_cols = [c for c in preview_cols if c in country_rows.columns]
            if not country_rows.empty:
                # ── Search bar ──────────────────────────────────────────────
                search_q = st.text_input(
                    "Search",
                    placeholder="Filter by title or year…",
                    key="editor_search",
                    label_visibility="collapsed",
                ).strip().lower()

                if search_q:
                    _mask = (
                        country_rows["Regulation_Title"].fillna("").str.lower().str.contains(search_q, regex=False)
                        | country_rows["Year"].astype(str).str.lower().str.contains(search_q, regex=False)
                    )
                    display_rows = country_rows[_mask].copy()
                else:
                    display_rows = country_rows

                st.caption("Expand a category, then click a row to edit or archive that record.")
                selected_entry_id = None

                if display_rows.empty and search_q:
                    st.info(f"No regulations match **{search_q}**.")
                else:
                    for cat in sorted(display_rows["Category"].dropna().unique().tolist()):
                        cat_rows = display_rows[display_rows["Category"] == cat].copy()
                        cat_preview = cat_rows[preview_cols].rename(columns={
                            "Regulation_Title": "Title",
                            "Source_URL": "Source URL",
                        }).copy()
                        if "Source URL" in cat_preview.columns:
                            cat_preview["Source URL"] = cat_preview["Source URL"].apply(
                                lambda u: safe_linkify(u) or None
                            )
                        cat_key = re.sub(r"[^a-zA-Z0-9]", "_", cat)
                        n = len(cat_rows)
                        label = f"{cat}  ·  {n} record{'s' if n != 1 else ''}"
                        with st.expander(label, expanded=bool(search_q)):
                            cat_event = st.dataframe(
                                cat_preview,
                                width="stretch",
                                hide_index=True,
                                height=min(300, max(80, 40 + n * 35)),
                                key=f"editor_table_{cat_key}",
                                on_select="rerun",
                                selection_mode="single-row",
                                column_config={
                                    "Source URL": st.column_config.LinkColumn(
                                        "Source URL",
                                        display_text="Open",
                                    ),
                                },
                            )
                            indices = get_selected_row_indices(cat_event)
                            if indices and selected_entry_id is None:
                                selected_entry_id = cat_rows.iloc[indices[0]]["Entry_ID"]

                if selected_entry_id:
                    existing_row = country_rows[
                        country_rows["Entry_ID"] == selected_entry_id
                    ].iloc[0]
                    edit_regulation_dialog(
                        existing_row=existing_row,
                        category_choices=category_choices,
                        df_all=df_all,
                        src_dir=resolved_data_dir,
                        editor_country=editor_country,
                        auth_user=str(_auth_user),
                    )

            else:
                st.info(
                    "No records exist yet for this country. "
                    "Use the form below to add a new regulation."
                )

            st.divider()
            if st.button("+ Add regulation", type="primary"):
                add_regulation_dialog(
                    category_choices=category_choices,
                    df_all=df_all,
                    src_dir=resolved_data_dir,
                    editor_country=editor_country,
                    auth_user=str(_auth_user),
                )
            st.divider()
            _ed_hist_df = pd.DataFrame()
            if ARCHIVE_FILE.exists():
                _ed_hist_df = pd.read_csv(ARCHIVE_FILE)
                _ed_hist_df["timestamp"] = pd.to_datetime(
                    _ed_hist_df["timestamp"], errors="coerce"
                )
                if not IS_ADMIN:
                    _ed_hist_df = _ed_hist_df[_ed_hist_df["country"] == editor_country]
                _ed_hist_df = _ed_hist_df.sort_values("timestamp", ascending=False).head(50)
                _ed_hist_df["changes"] = _ed_hist_df.apply(
                    lambda r: _audit_changes(
                        str(r.get("action", "")),
                        str(r.get("old_record", "") or ""),
                        str(r.get("new_record", "") or ""),
                    ),
                    axis=1,
                )
                _front = ["timestamp", "action", "country", "user", "changes"]
                _back  = [c for c in _ed_hist_df.columns if c not in _front]
                _ed_hist_df = _ed_hist_df[_front + _back]
            _ed_hist_col_cfg = {
                "timestamp": st.column_config.DatetimeColumn("Time", format="YYYY-MM-DD HH:mm"),
                "changes":   st.column_config.TextColumn("Changes", width="large"),
            }
            if st.button("📋 Change history", type="secondary", key="ed_hist_btn"):
                _dlg_change_history(_ed_hist_df, _ed_hist_col_cfg)

# Fallback: Gap Analysis or Editor accessed without authentication
elif _active_page in ("Gap Analysis", "Editor"):
    st.info("This page is available to authenticated users only. Contact the administrator for access.")


